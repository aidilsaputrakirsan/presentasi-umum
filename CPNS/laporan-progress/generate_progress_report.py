from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

def set_font(run, bold=False, size=11, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def shade_cell(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def add_placeholder_box(doc, kegiatan_num):
    """Menambahkan kotak placeholder untuk lampiran foto/video dokumentasi."""
    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run('Lampiran Dokumentasi Foto / Video')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_after = Pt(4)

    # Kotak placeholder sebagai tabel 1 baris
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(15)

    # Set background abu-abu muda
    shade_cell(cell, 'F2F3F4')

    # Konten dalam kotak
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run('[ AREA LAMPIRAN DOKUMENTASI KEGIATAN ' + str(kegiatan_num) + ' ]')
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(100, 100, 100)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Tempelkan foto / screenshot / tangkapan layar / video dokumentasi di area ini')
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(130, 130, 130)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('_' * 80)
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor(180, 180, 180)

    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run('Foto 1: ________________________________')
    r4.font.size = Pt(9)
    r4.font.color.rgb = RGBColor(130, 130, 130)

    p5 = cell.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run('Foto 2: ________________________________')
    r5.font.size = Pt(9)
    r5.font.color.rgb = RGBColor(130, 130, 130)

    p6 = cell.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r6 = p6.add_run('Video: ________________________________')
    r6.font.size = Pt(9)
    r6.font.color.rgb = RGBColor(130, 130, 130)

    p7 = cell.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r7 = p7.add_run('Keterangan: ________________________________')
    r7.font.size = Pt(9)
    r7.font.color.rgb = RGBColor(130, 130, 130)

    doc.add_paragraph()


# ====================== COVER PAGE ======================
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LAPORAN PROGRES AKTUALISASI')
set_font(run, bold=True, size=16, color=(0, 51, 102))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PELATIHAN DASAR CPNS ANGKATAN VI TAHUN 2026')
set_font(run, bold=True, size=13, color=(0, 51, 102))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SITRIA')
set_font(run, bold=True, size=22, color=(200, 146, 42))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sistem Informasi Tridharma Akademik')
set_font(run, bold=True, size=14, color=(26, 82, 118))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Dashboard Analitik Tridharma Dosen Berbasis Data Multi-Sumber')
set_font(run, italic=True, size=11, color=(89, 89, 89))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Lab Inovasi Digital - Prodi Sistem Informasi & Bisnis Digital')
set_font(run, size=11)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Fakultas Sains dan Teknologi Industri - Institut Teknologi Kalimantan')
set_font(run, size=11)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Disusun oleh:')
set_font(run, size=11, color=(89, 89, 89))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Aidil Saputra Kirsan')
set_font(run, bold=True, size=13, color=(0, 51, 102))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Dosen Sistem Informasi / Kepala Lab Inovasi Digital FSTI ITK')
set_font(run, size=11)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Tanggal Laporan: 7 April 2026')
set_font(run, bold=True, size=11, color=(200, 146, 42))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Periode Aktualisasi: 7 Maret - 22 April 2026 (Minggu 1-7)')
set_font(run, size=10, color=(89, 89, 89))

doc.add_page_break()

# ====================== RINGKASAN PROGRES ======================
h = doc.add_heading('RINGKASAN PROGRES AKTUALISASI', level=1)
h.runs[0].font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
run = p.add_run(
    'Per tanggal 7 April 2026 (Minggu ke-5 dari 7 minggu aktualisasi), berikut adalah ringkasan '
    'status penyelesaian setiap kegiatan dalam Rancangan Aktualisasi SITRIA:'
)
set_font(run, size=11)

doc.add_paragraph()

table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = table.rows[0].cells
headers = ['No', 'Kegiatan', 'Periode', 'Minggu', 'Status']
for cell, header in zip(hdr, headers):
    cell.text = header
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    shade_cell(cell, '003366')

kegiatan_summary = [
    ('1', 'Pemetaan & Pengumpulan Data Dosen Prodi SI & Bisnis Digital', '07-20 Mar', 'M1-M2', 'SELESAI (100%)', 'C6EFCE', '006100'),
    ('2', 'Pengembangan Sistem Pengambilan Data SINTA Otomatis', '14-27 Mar', 'M2-M3', 'SELESAI (100%)', 'C6EFCE', '006100'),
    ('3', 'Implementasi Fitur Analitik & Dashboard SITRIA', '21 Mar-10 Apr', 'M3-M5', 'HAMPIR SELESAI (~90%)', 'FFEB9C', '9C5700'),
    ('4', 'Implementasi Dashboard Akreditasi - DTPS & Pendanaan', '28 Mar-17 Apr', 'M4-M6', 'BERJALAN (~70%)', 'FFEB9C', '9C5700'),
    ('5', 'Peluncuran, Sosialisasi & Evaluasi Sistem SITRIA', '11-22 Apr', 'M6-M7', 'BELUM DIMULAI', 'F2DCDB', '9C0006'),
]

for row_data in kegiatan_summary:
    row = table.add_row()
    cells = row.cells
    num, name, period, week, status, bg, fg = row_data
    values = [num, name, period, week, status]
    for i, (cell, val) in enumerate(zip(cells, values)):
        cell.text = val
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
        para.runs[0].font.size = Pt(10)
        if i == 4:
            ri = int(fg[0:2], 16)
            gi = int(fg[2:4], 16)
            bi = int(fg[4:6], 16)
            para.runs[0].font.color.rgb = RGBColor(ri, gi, bi)
            para.runs[0].bold = True
            shade_cell(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

doc.add_paragraph()
doc.add_page_break()

# ====================== KEGIATAN DETAIL ======================

KEGIATAN = [
    {
        'num': 1,
        'title': 'Pemetaan & Pengumpulan Data Dosen Prodi SI & Bisnis Digital',
        'period': '07 - 20 Maret 2026',
        'week': 'Minggu 1-2',
        'status': 'SELESAI (100%)',
        'status_color': (0, 97, 0),
        'status_bg': 'C6EFCE',
        'agenda3': (
            'Kegiatan pemetaan dan pengumpulan data dosen merupakan bentuk nyata penerapan nilai Manajemen ASN, '
            'yaitu pengelolaan data kepegawaian secara digital, terstruktur, dan dapat dipertanggungjawabkan. '
            'Dengan membangun basis data dosen yang valid sebagai fondasi pengambilan keputusan, kegiatan ini '
            'sekaligus mencerminkan nilai Smart ASN dalam mendorong tata kelola akademik yang lebih efisien '
            'dan berbasis data di tingkat program studi.'
        ),
        'tahapan': [
            ('1', 'Koordinasi dengan Kaprodi SI dan Kaprodi Bisnis Digital', 'Notulensi rapat koordinasi', 'SELESAI'),
            ('2', 'Pengumpulan data SINTA ID seluruh dosen aktif', 'Daftar SINTA ID 2 prodi', 'SELESAI'),
            ('3', 'Verifikasi data melalui profil SINTA resmi', 'Dokumen verifikasi data dosen', 'SELESAI'),
            ('4', 'Penyusunan basis data dosen terstruktur', 'File basis data dosen terverifikasi', 'SELESAI'),
            ('5', 'Validasi bersama Mentor dan Kaprodi', 'Lembar validasi Kaprodi', 'SELESAI'),
        ],
        'capaian': [
            'Data seluruh dosen aktif dari 2 program studi (Sistem Informasi dan Bisnis Digital) berhasil '
            'dikumpulkan secara lengkap melalui koordinasi langsung dengan masing-masing Koordinator Program Studi.',
            'Setiap data dosen telah diverifikasi kebenarannya melalui profil resmi pada portal SINTA '
            'Kementerian Pendidikan, Kebudayaan, Riset, dan Teknologi, sehingga akurasi data terjamin.',
            'Basis data dosen tersusun dalam format digital terstruktur yang siap digunakan sebagai '
            'sumber data utama seluruh fitur analitik dalam sistem SITRIA.',
            'Proses dan hasil pemetaan telah mendapatkan validasi resmi dari Koordinator Program Studi '
            'masing-masing sebagai bentuk akuntabilitas kegiatan.',
        ],
        'catatan': (
            'Seluruh tahapan Kegiatan 1 telah diselesaikan tepat waktu pada minggu pertama dan kedua. '
            'Data dosen dari kedua program studi berhasil dikumpulkan, diverifikasi, dan disusun secara '
            'terstruktur sebagai fondasi sistem SITRIA. Capaian ini menjadi dasar yang kuat bagi '
            'keberhasilan pengembangan sistem pada kegiatan-kegiatan selanjutnya.'
        ),
    },
    {
        'num': 2,
        'title': 'Pengembangan Sistem Pengambilan Data SINTA Otomatis',
        'period': '14 - 27 Maret 2026',
        'week': 'Minggu 2-3',
        'status': 'SELESAI (100%)',
        'status_color': (0, 97, 0),
        'status_bg': 'C6EFCE',
        'agenda3': (
            'Pengembangan sistem pengambilan data otomatis merupakan wujud nyata nilai Smart ASN, '
            'di mana seorang ASN tidak sekadar menggunakan teknologi, melainkan aktif menciptakan solusi '
            'digital yang inovatif. Proses yang sebelumnya dilakukan secara manual dan membutuhkan waktu lama '
            'kini diubah menjadi sistem otomatis yang akuntabel, terukur, dan berkelanjutan, '
            'selaras dengan semangat transformasi digital dalam Manajemen ASN.'
        ),
        'tahapan': [
            ('1', 'Analisis struktur data pada portal SINTA Kemdiktisaintek', 'Dokumen analisis struktur data SINTA', 'SELESAI'),
            ('2', 'Pengembangan program pengambilan data otomatis (Python)', 'Program pengambil data SINTA aktif', 'SELESAI'),
            ('3', 'Implementasi penjadwalan otomatis (cron job)', 'Sistem penjadwalan otomatis aktif & terkonfigurasi', 'SELESAI'),
            ('4', 'Pengujian sistem dengan data 2 prodi', 'Data SINTA kedua prodi tersimpan & terstruktur', 'SELESAI'),
            ('5', 'Penyusunan dokumentasi dan panduan penggunaan', 'Dokumentasi teknis pipeline (README)', 'SELESAI'),
        ],
        'capaian': [
            'Sistem pengambilan data dari portal SINTA berhasil dikembangkan menggunakan bahasa pemrograman '
            'Python, mampu mengumpulkan data riset, publikasi, dan pengabdian seluruh dosen secara otomatis '
            'tanpa memerlukan intervensi manual.',
            'Data kinerja Tridharma seluruh dosen aktif dari kedua program studi berhasil terkumpul secara '
            'lengkap dan tersimpan dalam format terstruktur yang siap diolah oleh sistem SITRIA.',
            'Sistem penjadwalan otomatis telah dikonfigurasi sehingga pembaruan data dapat dilakukan '
            'secara berkala tanpa perlu dijalankan manual setiap saat.',
            'Proses yang sebelumnya memerlukan waktu berjam-jam untuk dilakukan secara manual kini '
            'dapat diselesaikan secara otomatis dalam hitungan menit.',
            'Dokumentasi lengkap panduan penggunaan sistem telah disusun untuk memastikan '
            'keberlanjutan pengoperasian oleh pengguna berikutnya.',
        ],
        'catatan': (
            'Seluruh tahapan Kegiatan 2 telah diselesaikan pada minggu kedua dan ketiga. '
            'Sistem pengambilan data otomatis berjalan penuh dan menghasilkan empat kelompok data '
            'per program studi: data statistik SINTA, data pengelompokan riset, data keahlian dosen, '
            'dan data peta perjalanan riset. Capaian ini menjadi tulang punggung seluruh fitur '
            'analitik yang dikembangkan pada Kegiatan 3 dan 4.'
        ),
    },
    {
        'num': 3,
        'title': 'Implementasi Fitur Analitik & Dashboard SITRIA',
        'period': '21 Maret - 10 April 2026',
        'week': 'Minggu 3-5',
        'status': 'HAMPIR SELESAI (~90%)',
        'status_color': (156, 87, 0),
        'status_bg': 'FFEB9C',
        'agenda3': (
            'Pembangunan fitur analitik dan dashboard merupakan puncak penerapan nilai Smart ASN dan '
            'Manajemen ASN dalam aktualisasi ini. Data yang semula tersebar dan sulit dibaca diubah '
            'menjadi informasi strategis yang mudah dipahami oleh pimpinan dan dosen. Pemanfaatan '
            'kecerdasan buatan untuk mengelompokkan topik riset secara otomatis menjadi bukti konkret '
            'inovasi digital ASN, sekaligus menyediakan instrumen pemantauan kinerja Tridharma yang '
            'akuntabel dan berbasis data.'
        ),
        'tahapan': [
            ('1', 'Implementasi pengelompokan topik riset otomatis (AI Clustering)', 'Fitur AI Clustering aktif & tervalidasi', 'SELESAI'),
            ('2', 'Pengembangan visualisasi perjalanan riset (Sankey Timeline)', 'Visualisasi peta riset terpasang', 'SELESAI'),
            ('3', 'Pembangunan Galeri Karya Akademik (6 kategori)', 'Galeri karya 6 kategori berfungsi penuh', 'SELESAI'),
            ('4', 'Penghubungan sistem tampilan dan sistem pengolahan data', 'Sistem terintegrasi & berjalan penuh', 'SELESAI'),
            ('5', 'Pengujian menyeluruh seluruh modul analitik', 'Laporan hasil pengujian modul', 'DALAM PROSES'),
        ],
        'capaian': [
            'Fitur pengelompokan topik riset berbasis kecerdasan buatan berhasil diimplementasikan dan '
            'mampu mendeteksi kesamaan topik penelitian antar dosen secara otomatis, termasuk '
            'potensi kolaborasi lintas program studi.',
            'Visualisasi peta perjalanan riset (Sankey Timeline) sudah dapat diakses melalui dashboard, '
            'menampilkan evolusi topik-topik penelitian dari tahun ke tahun sejak 2018 hingga saat ini.',
            'Galeri karya akademik dosen dengan 6 kategori (Penelitian, Pengabdian, Publikasi Scopus, '
            'SINTA, Buku, dan HKI) sudah tersedia dan dapat dicari serta difilter secara interaktif.',
            'Fitur pencarian pakar (Expertise Finder) sudah berfungsi, memudahkan mahasiswa atau '
            'mitra riset menemukan dosen yang paling sesuai dengan topik penelitian mereka.',
            'Seluruh komponen dashboard sudah terhubung dengan sistem data, sehingga informasi '
            'Tridharma dosen ditampilkan secara langsung dan selalu mutakhir.',
        ],
        'catatan': (
            '4 dari 5 tahapan Kegiatan 3 telah selesai diimplementasikan. Seluruh fitur utama '
            'dashboard analitik sudah berfungsi penuh dan dapat diakses. Satu tahapan yang masih '
            'berjalan adalah pengujian menyeluruh seluruh modul, yang ditargetkan selesai '
            'sebelum 10 April 2026 sesuai jadwal rancangan aktualisasi.'
        ),
    },
    {
        'num': 4,
        'title': 'Implementasi Dashboard Akreditasi - DTPS & Pendanaan Riset',
        'period': '28 Maret - 17 April 2026',
        'week': 'Minggu 4-6',
        'status': 'BERJALAN (~70%)',
        'status_color': (156, 87, 0),
        'status_bg': 'FFEB9C',
        'agenda3': (
            'Dashboard akreditasi dan pemantauan pendanaan riset merupakan bentuk konkret penerapan '
            'Manajemen ASN yang menekankan akuntabilitas kinerja berbasis data. Sistem kalkulasi '
            'otomatis rasio Dosen Tetap Program Studi (DTPS) sesuai standar BAN-PT menjamin '
            'objektivitas dan transparansi dalam pelaporan kinerja Tridharma. Kegiatan ini mendukung '
            'kesiapan akreditasi program studi secara nyata, sekaligus menjadi wujud Smart ASN '
            'yang berkontribusi pada peningkatan mutu institusi.'
        ),
        'tahapan': [
            ('1', 'Kajian standar indikator akreditasi program studi (LKPS BAN-PT)', 'Dokumen kajian standar LKPS', 'SELESAI'),
            ('2', 'Implementasi kalkulasi otomatis Rasio DTPS', 'Fitur kalkulasi DTPS otomatis aktif', 'SELESAI'),
            ('3', 'Pengembangan Dashboard Pemantauan Pendanaan Riset (BIMA)', 'Dashboard Pendanaan terpasang & berfungsi', 'SELESAI'),
            ('4', 'Integrasi data hibah eksternal dan internal kedua prodi', 'Data pendanaan riset terintegrasi penuh', 'SELESAI'),
            ('5', 'Validasi hasil kalkulasi DTPS bersama Wakil Dekan Akademik', 'Berita acara validasi Wakil Dekan', 'DIJADWALKAN'),
        ],
        'capaian': [
            'Dashboard akreditasi DTPS sudah menampilkan secara otomatis tiga indikator utama sesuai '
            'standar BAN-PT: rasio jumlah penelitian per dosen tetap, rasio pengabdian per dosen tetap, '
            'dan total dana riset per dosen tetap.',
            'Fitur simulasi interaktif tersedia bagi pimpinan untuk mensimulasikan dampak perubahan '
            'komposisi dosen (termasuk lintas prodi) terhadap nilai rasio akreditasi secara langsung.',
            'Dashboard pemantauan pendanaan riset sudah dapat menampilkan aliran dana hibah '
            'internal maupun eksternal (BIMA) dari masing-masing program studi secara terpisah maupun gabungan.',
            'Laporan progres triwulan pertama Lab Inovasi Digital sudah terdokumentasi dan tersedia '
            'sebagai rekam jejak perkembangan sistem.',
        ],
        'catatan': (
            '4 dari 5 tahapan Kegiatan 4 telah berjalan dengan baik. Fitur kalkulasi DTPS dan '
            'dashboard pendanaan riset sudah berfungsi penuh. Satu tahapan yang masih menunggu '
            'adalah validasi resmi bersama Wakil Dekan Akademik, yang segera dijadwalkan '
            'sebelum 17 April 2026.'
        ),
    },
    {
        'num': 5,
        'title': 'Peluncuran, Sosialisasi & Evaluasi Sistem SITRIA',
        'period': '11 - 22 April 2026',
        'week': 'Minggu 6-7',
        'status': 'BELUM DIMULAI',
        'status_color': (156, 0, 6),
        'status_bg': 'F2DCDB',
        'agenda3': (
            'Peluncuran dan sosialisasi sistem merupakan tahap krusial dalam mewujudkan nilai Smart ASN '
            'yang sesungguhnya: tidak hanya membangun sistem, tetapi juga membangun budaya digital '
            'di lingkungan unit kerja. Proses transfer pengetahuan kepada seluruh pengguna mendukung '
            'nilai Manajemen ASN dalam menjamin keberlanjutan dan kemandirian pengelolaan sistem. '
            'Evaluasi berbasis masukan pengguna mencerminkan ASN yang adaptif dan responsif '
            'terhadap kebutuhan nyata di lapangan.'
        ),
        'tahapan': [
            ('1', 'Peluncuran sistem SITRIA ke server Lab Inovasi Digital', 'Sistem SITRIA aktif & dapat diakses online', 'DIJADWALKAN 11 APR'),
            ('2', 'Penyusunan Panduan Pengguna & SOP pengelolaan data', 'Dokumen Panduan Pengguna & SOP tersusun', 'DIJADWALKAN 11-13 APR'),
            ('3', 'Sosialisasi dan pelatihan pengguna (dosen, kaprodi, laboran)', 'Daftar hadir & dokumentasi kegiatan sosialisasi', 'DIJADWALKAN 14-16 APR'),
            ('4', 'Pengumpulan masukan & evaluasi sistem dari pengguna', 'Laporan evaluasi & rekapitulasi masukan', 'DIJADWALKAN 17-19 APR'),
            ('5', 'Penyusunan laporan aktualisasi akhir', 'Laporan aktualisasi final', 'DIJADWALKAN 20-22 APR'),
        ],
        'capaian': [
            'Kegiatan ini belum dimulai sesuai jadwal yang ditetapkan, akan dilaksanakan '
            'mulai 11 April 2026.',
            'Seluruh fitur sistem SITRIA yang dibangun pada Kegiatan 1 hingga 4 sudah dalam '
            'kondisi siap untuk diluncurkan ke server.',
            'Materi sosialisasi dan bahan pelatihan pengguna sedang dalam tahap persiapan awal.',
        ],
        'catatan': (
            'Kegiatan 5 belum dimulai sesuai jadwal rancangan aktualisasi. Namun demikian, '
            'seluruh fondasi teknis sistem SITRIA yang dibangun pada Kegiatan 1-4 sudah siap '
            'untuk mendukung proses peluncuran. Target akhir: sistem SITRIA aktif, '
            'pengguna terlatih, dan laporan aktualisasi tersusun sebelum 22 April 2026.'
        ),
    },
]

for k in KEGIATAN:
    h = doc.add_heading('KEGIATAN ' + str(k['num']) + ': ' + k['title'], level=1)
    h.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    h.runs[0].font.size = Pt(13)

    # Info bar
    info_table = doc.add_table(rows=1, cols=3)
    info_table.style = 'Table Grid'
    info_cells = info_table.rows[0].cells
    info_data = [('Periode', k['period']), ('Minggu', k['week']), ('Status', k['status'])]
    colors = ['EBF5FB', 'EBF5FB', k['status_bg']]
    for i, ((label, val), bg) in enumerate(zip(info_data, colors)):
        cell = info_cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + '\n')
        r1.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = RGBColor(0, 51, 102)
        r2 = p.add_run(val)
        r2.bold = True
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(*k['status_color']) if i == 2 else RGBColor(0, 51, 102)
        shade_cell(cell, bg)

    doc.add_paragraph()

    # Agenda III
    p = doc.add_paragraph()
    r = p.add_run('Korelasi dengan Agenda III - Kedudukan dan Peran PNS dalam NKRI')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_after = Pt(3)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(k['agenda3'])
    r.font.size = Pt(10)
    r.italic = True

    # Tahapan table
    p = doc.add_paragraph()
    r = p.add_run('Tahapan Kegiatan & Output')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_after = Pt(3)

    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    hrow = t.rows[0].cells
    for cell, txt in zip(hrow, ['No', 'Tahapan Kegiatan', 'Output / Bukti Fisik', 'Status']):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(cell, '1A5276')

    for no, tahap, output, status in k['tahapan']:
        row = t.add_row()
        cells = row.cells
        vals = [no, tahap, output, status]
        aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER]
        for cell, val, align in zip(cells, vals, aligns):
            cell.text = val
            cell.paragraphs[0].alignment = align
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if 'SELESAI' in val and 'HAMPIR' not in val and 'DIJADWALKAN' not in val and 'PROSES' not in val:
                run.font.color.rgb = RGBColor(0, 97, 0)
                run.bold = True
                shade_cell(cell, 'EAFAF1')
            elif 'DIJADWALKAN' in val or 'PROSES' in val:
                run.font.color.rgb = RGBColor(156, 87, 0)
                run.bold = True

    doc.add_paragraph()

    # Capaian (menggantikan "Bukti Kemajuan dari Codebase")
    p = doc.add_paragraph()
    r = p.add_run('Capaian & Hasil Nyata Kegiatan')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_after = Pt(3)

    for cap in k['capaian']:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(cap)
        r.font.size = Pt(10)

    doc.add_paragraph()

    # Catatan
    p = doc.add_paragraph()
    r = p.add_run('Catatan Progres')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_after = Pt(3)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(k['catatan'])
    r.font.size = Pt(10)

    # Placeholder lampiran foto/video
    add_placeholder_box(doc, k['num'])

    if k['num'] < 5:
        doc.add_page_break()

# ====================== RENCANA TINDAK LANJUT ======================
doc.add_page_break()

h = doc.add_heading('RENCANA TINDAK LANJUT', level=1)
h.runs[0].font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
r = p.add_run(
    'Berdasarkan progres per 7 April 2026, berikut adalah rencana tindak lanjut yang akan segera dikerjakan '
    'dalam sisa periode aktualisasi (8 - 22 April 2026):'
)
r.font.size = Pt(11)

doc.add_paragraph()

rtl_data = [
    ('8 - 10 Apr 2026', 'Menyelesaikan pengujian menyeluruh seluruh modul analitik (Kegiatan 3 - Tahap 5)', 'Tinggi'),
    ('10 - 13 Apr 2026', 'Koordinasi dan pelaksanaan validasi DTPS bersama Wakil Dekan Akademik (Kegiatan 4 - Tahap 5)', 'Tinggi'),
    ('11 Apr 2026', 'Peluncuran sistem SITRIA ke server Lab Inovasi Digital (Kegiatan 5 - Tahap 1)', 'Tinggi'),
    ('11 - 13 Apr 2026', 'Penyusunan Panduan Pengguna & SOP pengelolaan data (Kegiatan 5 - Tahap 2)', 'Sedang'),
    ('14 - 16 Apr 2026', 'Sosialisasi dan pelatihan pengguna: dosen, Kaprodi, dan laboran (Kegiatan 5 - Tahap 3)', 'Tinggi'),
    ('17 - 19 Apr 2026', 'Pengumpulan masukan pengguna dan evaluasi sistem (Kegiatan 5 - Tahap 4)', 'Sedang'),
    ('20 - 22 Apr 2026', 'Finalisasi dan penyelesaian laporan aktualisasi akhir (Kegiatan 5 - Tahap 5)', 'Tinggi'),
]

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
for cell, txt in zip(t.rows[0].cells, ['Tanggal', 'Rencana Kegiatan', 'Prioritas']):
    cell.text = txt
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_cell(cell, '003366')

for date, act, prio in rtl_data:
    row = t.add_row()
    cells = row.cells
    cells[0].text = date
    cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    cells[1].text = act
    cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    cells[2].text = prio
    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cells[2].paragraphs[0].runs[0]
    r.bold = True
    r.font.size = Pt(10)
    if prio == 'Tinggi':
        r.font.color.rgb = RGBColor(156, 0, 6)
        shade_cell(cells[2], 'F2DCDB')
    else:
        r.font.color.rgb = RGBColor(156, 87, 0)
        shade_cell(cells[2], 'FFEB9C')

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('Target: Seluruh 5 Kegiatan Aktualisasi SITRIA selesai 100% pada 22 April 2026.')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0, 51, 102)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

output_path = r'D:\Github-ADL\presentasi-umum\CPNS\laporan-progress\Laporan_Progress_Aktualisasi_SITRIA_7Apr2026_v2.docx'
doc.save(output_path)
print('Saved: ' + output_path)
