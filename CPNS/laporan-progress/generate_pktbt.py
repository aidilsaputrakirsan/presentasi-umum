from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def set_run(run, size=9, bold=False, name='Times New Roman'):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold


def set_col_width(row_obj, widths):
    for j, w in enumerate(widths):
        row_obj.cells[j].width = Cm(w)


# ============================================================
doc = Document()

for sec in doc.sections:
    sec.page_height = Cm(21.0)
    sec.page_width = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.0)

normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(11)

# ---- TITLE ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('REKAPITULASI NILAI PENGUATAN KOMPETENSI TEKNIK SUBSTANSTIF')
set_run(r, size=12, bold=True)
doc.add_paragraph()

# ---- IDENTITY ----
id_items = [
    ('Program', 'Pelatihan Dasar CPNS'),
    ('Nama Peserta', 'Aidil Saputra Kirsan, S.ST., M.Tr.Kom'),
    ('NIP', '199403172025061004'),
    ('Jabatan', 'Penata Muda Tk. I (III/b)'),
]
for label, value in id_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    t = OxmlElement('w:tab')
    t.set(qn('w:val'), 'left')
    t.set(qn('w:pos'), '2000')
    tabs.append(t)
    pPr.append(tabs)
    r = p.add_run(label)
    set_run(r, size=11)
    r2 = p.add_run('\t: ')
    set_run(r2, size=11)
    r3 = p.add_run(value)
    set_run(r3, size=11)

doc.add_paragraph()

# ---- MAIN TABLE DATA ----
PKTBT_DATA = [
    dict(
        standar='Sertifikasi Kompetensi Junior Web Developer',
        jenis='Klasikal',
        tujuan='Memperkuat kompetensi ASN dalam pengembangan aplikasi web front-end berbasis standar nasional BNSP guna mendukung tugas pengajaran dan pembimbingan mahasiswa di bidang pemrograman web.',
        strategi='Ujian Kompetensi',
        mata=[
            'Membuat halaman web menggunakan HTML5 dan CSS3',
            'Mengimplementasikan interaktivitas dengan JavaScript',
            'Mengembangkan tampilan web responsif (Responsive Web Design)',
            'Melakukan pengujian dan debugging aplikasi web',
        ],
        jp='8 JP',
        tempat='Balikpapan, melalui LSP Telematika Indonesia',
    ),
    dict(
        standar='Sertifikasi Kompetensi Auditor Teknologi Informasi (CITA - Certified IT Auditor)',
        jenis='Klasikal',
        tujuan='Memperkuat kapasitas ASN dalam melaksanakan audit teknologi informasi, tata kelola TI, dan manajemen risiko sistem informasi sesuai standar internasional.',
        strategi='Pelatihan/Bimbingan Teknis dan Ujian Kompetensi',
        mata=[
            'Tata Kelola dan Manajemen Teknologi Informasi',
            'Framework Audit TI (COBIT, ISO/IEC 27001)',
            'Perencanaan dan Pelaksanaan Audit TI',
            'Manajemen Risiko dan Keamanan Informasi',
            'Pelaporan Hasil Audit Teknologi Informasi',
        ],
        jp='40 JP',
        tempat='Yogyakarta, diselenggarakan oleh TUV Rheinland Indonesia',
    ),
    dict(
        standar='Pelatihan Instruktur/Pengajar Junior Web Developer BNSP',
        jenis='Klasikal',
        tujuan='Memperkuat kompetensi ASN sebagai pengajar pelatihan berbasis kompetensi BNSP di bidang pengembangan web, sehingga mampu melatih dan mentransfer keterampilan teknis kepada peserta didik secara efektif.',
        strategi='Pelatihan/Bimbingan Teknis',
        mata=[
            'Metodologi Pelatihan Berbasis Kompetensi (PBK)',
            'Pengembangan Materi Ajar Junior Web Developer',
            'Teknik Fasilitasi dan Penyampaian Materi Teknis',
            'Penilaian dan Evaluasi Peserta Pelatihan',
        ],
        jp='24 JP',
        tempat='Balikpapan',
    ),
    dict(
        standar='Pelatihan Digital Talent Scholarship - Junior Networking',
        jenis='Non-Klasikal',
        tujuan='Menambah penguatan literasi dan kompetensi teknis jaringan komputer bagi ASN untuk mendukung peran dalam transformasi digital dan pengelolaan infrastruktur TI di lingkungan kerja.',
        strategi='Pelatihan/Bimbingan Teknis',
        mata=[
            'Dasar-dasar Jaringan Komputer',
            'Konfigurasi dan Manajemen Perangkat Jaringan',
            'Keamanan Jaringan Komputer',
            'Troubleshooting Jaringan',
            'Pengantar Jaringan Berbasis Cloud',
        ],
        jp='40 JP',
        tempat='Daring melalui platform Digital Talent Scholarship, Kementerian Komunikasi dan Informatika',
    ),
    dict(
        standar='E-Learning Pemrograman JavaScript (Dicoding Indonesia)',
        jenis='Non-Klasikal',
        tujuan='Memperkuat kompetensi teknis ASN dalam pemrograman berbasis JavaScript sebagai pondasi pengembangan aplikasi web modern yang mendukung tugas pengajaran dan pengembangan sistem digital.',
        strategi='E-Learning',
        mata=[
            'Dasar-dasar Pemrograman JavaScript',
            'Tipe Data, Variabel, dan Operator',
            'Struktur Kontrol dan Fungsi',
            'Array, Objek, dan Pemrograman Berbasis Prototipe',
            'DOM Manipulation dan Event Handling',
        ],
        jp='20 JP',
        tempat='E-Learning Dicoding Indonesia (dicoding.com)',
    ),
    dict(
        standar='Literasi Digital',
        jenis='Non-Klasikal',
        tujuan='Memberikan penguatan literasi digital bagi ASN serta adaptif dan bijak dalam menggunakan teknologi digital di era transformasi layanan publik.',
        strategi='E-Learning pengetahuan literasi digital',
        mata=[
            'Etika Digital',
            'Budaya Digital',
            'Keamanan Digital',
            'Kecakapan Digital',
        ],
        jp='14 JP',
        tempat='E-Learning',
    ),
    dict(
        standar='Pengetahuan Antikorupsi Dasar dan Integritas (PADI)',
        jenis='Non-Klasikal',
        tujuan='Memberikan pemahaman serta mencegah ASN dari perbuatan korupsi yang merugikan diri sendiri dan sekitar, serta membangun integritas ASN yang kuat.',
        strategi='E-Learning pengetahuan antikorupsi dasar dan integritas (PADI) untuk umum',
        mata=[
            'Mengapa korupsi harus diberantas?',
            'Pahami dulu baru lawan',
            'Pilih strategi, rancang aksi',
            'Aksi Integritas berantas korupsi',
            'Bangun kompetensi dengan sertifikasi antikorupsi',
        ],
        jp='20 JP',
        tempat='E-Learning',
    ),
    dict(
        standar='Organisasi dan Tata Kerja Institut Teknologi Kalimantan',
        jenis='Non-Klasikal',
        tujuan='Memberikan penguatan pemahaman ASN akan tata kelola dan struktur organisasi Institut Teknologi Kalimantan sebagai landasan pelaksanaan tugas dan fungsi jabatan.',
        strategi='Penguatan tata kelola',
        mata=[
            'Materi kelembagaan institusi negara',
        ],
        jp='1 JP',
        tempat='Zoom Meeting',
    ),
]

n_data = len(PKTBT_DATA)
n_rows = 1 + n_data + 2
n_cols = 9

# Column widths (cm) - total = 25.2 cm (landscape A4: 29.7 - 2.5 - 2.0 = 25.2)
CW = [0.6, 3.5, 1.8, 3.8, 2.8, 5.0, 1.3, 3.5, 0.9]

tbl = doc.add_table(rows=n_rows, cols=n_cols)
tbl.style = 'Table Grid'

# ---- HEADER ROW ----
hdr = tbl.rows[0]
trPr = hdr._tr.get_or_add_trPr()
tblHeader = OxmlElement('w:tblHeader')
trPr.append(tblHeader)

HDR_LABELS = [
    'No.',
    'Standar\nKompetensi',
    'Jenis\nPenguatan\nKompetensi',
    'Tujuan\nPenguatan',
    'Strategi/\nMetode\nPenguatan',
    'Mata pelatihan',
    'Jumlah\nJP/\nHari',
    'Tempat\nPelaksanaan',
    'Nilai',
]
set_col_width(hdr, CW)
for j, label in enumerate(HDR_LABELS):
    c = hdr.cells[j]
    set_cell_bg(c, 'D9D9D9')
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label)
    set_run(r, size=9, bold=True)

# ---- DATA ROWS ----
for i, item in enumerate(PKTBT_DATA):
    row = tbl.rows[i + 1]
    set_col_width(row, CW)

    # Col 0: No
    c = row.cells[0]
    c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(f'{i + 1}.'), size=9)

    # Col 1: Standar Kompetensi
    c = row.cells[1]
    c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_run(c.paragraphs[0].add_run(item['standar']), size=9)

    # Col 2: Jenis
    c = row.cells[2]
    c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(item['jenis']), size=9)

    # Col 3: Tujuan
    c = row.cells[3]
    c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_run(c.paragraphs[0].add_run(item['tujuan']), size=9)

    # Col 4: Strategi
    c = row.cells[4]
    c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_run(c.paragraphs[0].add_run(item['strategi']), size=9)

    # Col 5: Mata Pelatihan
    c = row.cells[5]
    c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for k, mp in enumerate(item['mata']):
        p = c.paragraphs[0] if k == 0 else c.add_paragraph()
        set_run(p.add_run(f'{k + 1}. {mp}'), size=9)

    # Col 6: JP
    c = row.cells[6]
    c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(item['jp']), size=9)

    # Col 7: Tempat
    c = row.cells[7]
    c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_run(c.paragraphs[0].add_run(item['tempat']), size=9)

    # Col 8: Nilai (empty)
    c = row.cells[8]
    c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

# ---- TOTAL ROWS ----
for k, label_text in enumerate(['NILAI TOTAL (RATA-RATA)', 'NILAI AKHIR KTS (60%)']):
    row = tbl.rows[n_data + 1 + k]
    set_col_width(row, CW)
    merged = row.cells[0].merge(row.cells[7])
    set_cell_bg(merged, 'D9D9D9')
    merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = merged.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(label_text), size=9, bold=True)
    set_cell_bg(row.cells[8], 'D9D9D9')

# ---- SIGNATURE SECTION ----
doc.add_paragraph()

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(p_date.add_run('Balikpapan, 04 April 2026'), size=11)

doc.add_paragraph()

# Mentor only - right aligned block
p_mentor_label = doc.add_paragraph()
p_mentor_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(p_mentor_label.add_run('Mentor'), size=11)
for _ in range(4):
    p = doc.add_paragraph()
    set_run(p.add_run(''), size=11)
p_mn = doc.add_paragraph()
p_mn.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(p_mn.add_run('Irma Fitria, S.Si., M.Si.'), size=11, bold=True)
p_mnip = doc.add_paragraph()
p_mnip.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(p_mnip.add_run('NIP. 199303232022032016'), size=11)

# ---- SAVE ----
out = 'c:/laragon/www/Materi-Presentasi/presentasi-umum/CPNS/laporan-progress/Aidil_Saputra_Kirsan_PKTBT_199403172025061004.docx'
doc.save(out)
print(f'SUCCESS: {out}')
