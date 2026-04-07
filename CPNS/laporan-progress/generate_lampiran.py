"""
Generator: Lampiran Bukti Aktualisasi SITRIA
Menghasilkan: Lampiran_Bukti_Aktualisasi_SITRIA.docx
"""
import json
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# DATA AKTUAL DARI JSON
# ============================================================
SI_LECTURERS = [
    ("Yuyun Tri Wiranti",                        "5978281", 6,  40, 8,  9,  5),
    ("Aidil Saputra Kirsan",                      "6760340", 6,  28, 5,  9,  6),
    ("Arif Wicaksono Septyanto",                  "6741019", 6,  23, 5,  9,  5),
    ("Henokh Lugo Hariyanto",                     "6807418", 0,  9,  3,  6,  2),
    ("Lovinta Happy Atrinawati",                  "5977894", 0,  53, 12, 3,  1),
    ("Vika Fitratunnany Insanittaqwa",            "6784235", 0,  14, 3,  5,  3),
    ("Hendy Indrawan Sunardi",                    "6784468", 0,  9,  2,  2,  5),
    ("Dwi Nur Amalia",                            "6761367", 0,  14, 3,  3,  4),
    ("Dwi Arief Prambudi",                        "6784358", 0,  32, 7,  7,  7),
    ("I Putu Deny Arthawan Sugih Prabowo",        "6701825", 4,  42, 6,  8,  5),
    ("M. Ihsan Alfani Putera",                    "6681873", 3,  40, 7,  7,  5),
    ("Sri Rahayu Natasia",                        "5983406", 10, 56, 9,  10, 5),
    ("Nursanti Novi Arisa",                       "6757976", 0,  18, 6,  6,  6),
    ("Rosa Eliviani",                             "6876629", 0,  13, 3,  1,  1),
    ("M. Gilvy Langgawan Putra",                  "6201079", 0,  70, 12, 1,  2),
]

BD_LECTURERS = [
    ("Agung Prabowo, S.E., M.M.",                           "6803848", 0, 14, 3, 3,  2),
    ("Bayu Nur Abdallah, S.T., M.T., CSCA.",               "6686868", 5, 27, 5, 10, 7),
    ("Deli Yansyah, S.E., M.Acc., Ak.",                    "6807445", 0, 8,  2, 4,  3),
    ("Eka Krisna Santoso, S.Si., MBA.",                     "6788207", 0, 15, 3, 0,  3),
    ("Fegy Sukris Sri Andriany, M.Ak.",                     "6957012", 0, 0,  0, 0,  0),
    ("Ir. Riovan Styx Roring, S.T., M.Kom.",               "6958596", 1, 20, 6, 3,  1),
    ("Khairunnisa Rahmah, S.E., M.M., CDMS.",              "6798949", 0, 24, 6, 5,  3),
    ("Luh Made Wisnu Satyaninggrat, S.Kom., M.T.",         "6793011", 1, 26, 4, 9,  5),
    ("Muhammad Ikhsan Alif S., S.E., M.Sc., CPEC., CHCM.","6780638", 3, 20, 3, 5,  7),
    ("Prasis Damai Nursyam Hamijaya, S.P., M.M.",          "6790854", 0, 25, 3, 7,  3),
]

# ============================================================
# HELPERS
# ============================================================
def shade_cell(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell_border_all(cell, color='CCCCCC', sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        elm = OxmlElement(f'w:{edge}')
        elm.set(qn('w:val'), 'single')
        elm.set(qn('w:sz'), str(sz))
        elm.set(qn('w:space'), '0')
        elm.set(qn('w:color'), color)
        borders.append(elm)
    tcPr.append(borders)

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    col = (0, 51, 102) if level == 1 else (26, 82, 118)
    h.runs[0].font.color.rgb = RGBColor(*col)
    h.runs[0].font.size = Pt(13 if level == 1 else 11)
    return h

def subheading(doc, text, color=(0, 51, 102)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(4)
    return p

def normal(doc, text, indent=0, italic=False, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    return p

def tbl_header(tbl, headers, bg='003366'):
    row = tbl.rows[0].cells
    for cell, txt in zip(row, headers):
        cell.text = txt
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        shade_cell(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_row(tbl, values, aligns=None, bold_cols=None, bg=None):
    row = tbl.add_row()
    if aligns is None:
        aligns = [WD_ALIGN_PARAGRAPH.CENTER] * len(values)
    if bold_cols is None:
        bold_cols = []
    for i, (cell, val) in enumerate(zip(row.cells, values)):
        cell.text = str(val)
        cell.paragraphs[0].alignment = aligns[i]
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        run.bold = i in bold_cols
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if bg:
            shade_cell(cell, bg)

def screenshot_box(doc, filename, caption, kegiatan_num, tahap_num):
    """Kotak placeholder screenshot dengan nama file yang spesifik."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f'Lampiran {kegiatan_num}.{tahap_num} — {caption}')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_after = Pt(3)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, 'EBF5FB')

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run('[ TEMPELKAN SCREENSHOT DI SINI ]')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(26, 82, 118)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f'Nama File: {filename}')
    r2.bold = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(200, 146, 42)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(caption)
    r3.italic = True
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor(100, 100, 100)

    # Spacer rows
    for _ in range(5):
        cell.add_paragraph()

    doc.add_paragraph()

# ============================================================
# BUILD DOCUMENT
# ============================================================
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

# ---- COVER ----
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('LAMPIRAN BUKTI AKTUALISASI')
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SISTEM INFORMASI TRIDHARMA AKADEMIK (SITRIA)')
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Lab Inovasi Digital - FSTI Institut Teknologi Kalimantan')
r.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Aidil Saputra Kirsan | Dosen SI / Kepala Lab Inovasi Digital')
r.font.size = Pt(11); r.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Dokumen ini berisi bukti fisik dan output dari setiap tahapan kegiatan aktualisasi.')
r.italic = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(89, 89, 89)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Periode Aktualisasi: 7 Maret - 22 April 2026')
r.font.size = Pt(10); r.bold = True; r.font.color.rgb = RGBColor(200, 146, 42)

doc.add_paragraph()
normal(doc, 'Keterangan pengisian lampiran:', size=9)
items = [
    'Tabel data: diisi otomatis dari data sistem SITRIA yang telah berjalan.',
    'Kotak biru [TEMPELKAN SCREENSHOT]: ambil tangkapan layar sesuai nama file yang tertera, lalu tempelkan gambar ke dalam kotak tersebut.',
    'Dokumen template: isi/lengkapi bagian yang diberi tanda garis bawah (___) sesuai kondisi nyata.',
]
for item in items:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(item)
    r.font.size = Pt(9)

doc.add_page_break()


# ============================================================
# LAMPIRAN KEGIATAN 1
# ============================================================
heading(doc, 'LAMPIRAN KEGIATAN 1')
normal(doc, 'Pemetaan & Pengumpulan Data Dosen Prodi SI & Bisnis Digital | 07 - 20 Maret 2026', italic=True)
doc.add_paragraph()

# --- L1.1 Notulensi Rapat Koordinasi ---
subheading(doc, 'Lampiran 1.1 — Notulensi Rapat Koordinasi dengan Kaprodi')

# Box notulensi
tbl = doc.add_table(rows=1, cols=1)
tbl.style = 'Table Grid'
cell = tbl.rows[0].cells[0]
shade_cell(cell, 'FDFEFE')

lines = [
    ('NOTULENSI RAPAT KOORDINASI', True, 11, (0, 51, 102)),
    ('Pengembangan Sistem Informasi Tridharma Akademik (SITRIA)', True, 10, (0, 51, 102)),
    ('', False, 9, (0,0,0)),
    ('Hari/Tanggal  : ___________________________ (Maret 2026)', False, 9, (50,50,50)),
    ('Waktu          : ___________________________', False, 9, (50,50,50)),
    ('Tempat         : ___________________________', False, 9, (50,50,50)),
    ('Dipimpin oleh  : ___________________________', False, 9, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('Peserta Rapat:', True, 9, (0, 51, 102)),
    ('  1. Aidil Saputra Kirsan — Dosen SI / Kepala Lab Inovasi Digital (Pemrakarsa)', False, 9, (50,50,50)),
    ('  2. ___________________________ — Koordinator Prodi Sistem Informasi', False, 9, (50,50,50)),
    ('  3. ___________________________ — Koordinator Prodi Bisnis Digital', False, 9, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('Agenda Rapat:', True, 9, (0, 51, 102)),
    ('  1. Pemaparan rancangan sistem SITRIA dan tujuan aktualisasi Latsar CPNS', False, 9, (50,50,50)),
    ('  2. Koordinasi pengumpulan data SINTA ID dosen aktif masing-masing prodi', False, 9, (50,50,50)),
    ('  3. Pembahasan kebutuhan data dan fitur dashboard prioritas', False, 9, (50,50,50)),
    ('  4. Penetapan jadwal dan mekanisme validasi data', False, 9, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('Hasil & Kesepakatan Rapat:', True, 9, (0, 51, 102)),
    ('  1. Kaprodi SI dan Bisnis Digital menyetujui dan mendukung pengembangan sistem SITRIA.', False, 9, (50,50,50)),
    ('  2. Pengumpulan data SINTA ID seluruh dosen aktif disepakati dalam waktu 3 hari kerja.', False, 9, (50,50,50)),
    ('  3. Fitur prioritas yang disepakati: Dashboard SINTA, Akreditasi DTPS, dan Funding.', False, 9, (50,50,50)),
    ('  4. Validasi data dilakukan bersama Kaprodi sebelum dimasukkan ke sistem.', False, 9, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('Tindak Lanjut:', True, 9, (0, 51, 102)),
    ('  - Aidil: mengumpulkan SINTA ID dari masing-masing prodi', False, 9, (50,50,50)),
    ('  - Kaprodi SI: menyiapkan daftar dosen aktif Prodi SI', False, 9, (50,50,50)),
    ('  - Kaprodi BD: menyiapkan daftar dosen aktif Prodi Bisnis Digital', False, 9, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('Mengetahui,', True, 9, (0, 51, 102)),
    ('', False, 9, (0,0,0)),
    ('Koordinator Prodi SI                  Koordinator Prodi BD                  Pemrakarsa', False, 9, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('', False, 9, (0,0,0)),
    ('', False, 9, (0,0,0)),
    ('(________________________)           (________________________)           (Aidil Saputra Kirsan)', False, 9, (50,50,50)),
]
first = True
for text, bold, size, color in lines:
    if first:
        p = cell.paragraphs[0]
        first = False
    else:
        p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold and size >= 10 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*color)

doc.add_paragraph()

# Screenshot koordinasi
screenshot_box(doc, 'bukti_k1_01_foto_koordinasi.jpg',
               'Foto dokumentasi saat rapat koordinasi dengan Kaprodi SI dan Bisnis Digital',
               1, '1b')

doc.add_page_break()

# --- L1.2 Daftar SINTA ID ---
subheading(doc, 'Lampiran 1.2 — Daftar SINTA ID Dosen Aktif (2 Program Studi)')
normal(doc, 'Sumber data: sistem SITRIA | Diverifikasi melalui portal sinta.kemdiktisaintek.go.id', italic=True)
doc.add_paragraph()

# Tabel SI
subheading(doc, 'A. Prodi Sistem Informasi (15 Dosen)', color=(26, 82, 118))
tbl_si = doc.add_table(rows=1, cols=5)
tbl_si.style = 'Table Grid'
tbl_header(tbl_si, ['No', 'Nama Dosen', 'SINTA ID', 'URL Profil SINTA', 'Status Verifikasi'])
for i, (name, sid, sc, gs, h, res, pkm) in enumerate(SI_LECTURERS, 1):
    url = f'sinta.kemdiktisaintek.go.id/authors/profile/{sid}'
    add_row(tbl_si,
            [i, name, sid, url, 'Terverifikasi'],
            aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                    WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                    WD_ALIGN_PARAGRAPH.CENTER],
            bold_cols=[],
            bg='EAFAF1' if i % 2 == 0 else None)
    # Color status
    row = tbl_si.rows[i]
    row.cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 97, 0)
    row.cells[4].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Tabel BD
subheading(doc, 'B. Prodi Bisnis Digital (10 Dosen)', color=(26, 82, 118))
tbl_bd = doc.add_table(rows=1, cols=5)
tbl_bd.style = 'Table Grid'
tbl_header(tbl_bd, ['No', 'Nama Dosen', 'SINTA ID', 'URL Profil SINTA', 'Status Verifikasi'])
for i, (name, sid, sc, gs, h, res, pkm) in enumerate(BD_LECTURERS, 1):
    url = f'sinta.kemdiktisaintek.go.id/authors/profile/{sid}'
    add_row(tbl_bd,
            [i, name, sid, url, 'Terverifikasi'],
            aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                    WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                    WD_ALIGN_PARAGRAPH.CENTER],
            bg='EBF5FB' if i % 2 == 0 else None)
    row = tbl_bd.rows[i]
    row.cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 97, 0)
    row.cells[4].paragraphs[0].runs[0].bold = True

doc.add_paragraph()
screenshot_box(doc, 'bukti_k1_02_sinta_profil_dosen.png',
               'Tangkapan layar contoh profil dosen di portal SINTA (sinta.kemdiktisaintek.go.id)',
               1, '2b')

doc.add_page_break()

# --- L1.3 Dokumen Verifikasi Data Dosen ---
subheading(doc, 'Lampiran 1.3 — Dokumen Verifikasi Data Dosen')
normal(doc, 'Verifikasi dilakukan dengan mencocokkan nama dosen pada portal SINTA resmi menggunakan SINTA ID yang telah dikumpulkan.', italic=True)
doc.add_paragraph()

tbl_v = doc.add_table(rows=1, cols=6)
tbl_v.style = 'Table Grid'
tbl_header(tbl_v, ['No', 'Nama Dosen', 'Prodi', 'SINTA ID', 'Nama di SINTA Sesuai', 'Profil Aktif'])

all_lecs = [(n, s, 'SI') for n, s, *_ in SI_LECTURERS] + [(n, s, 'BD') for n, s, *_ in BD_LECTURERS]
for i, (name, sid, prodi) in enumerate(all_lecs, 1):
    row = tbl_v.add_row()
    cells = row.cells
    vals = [i, name, prodi, sid, 'Ya', 'Ya']
    aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
              WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
    for j, (cell, val, align) in enumerate(zip(cells, vals, aligns)):
        cell.text = str(val)
        cell.paragraphs[0].alignment = align
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        if j in (4, 5):
            run.font.color.rgb = RGBColor(0, 97, 0)
            run.bold = True
        if i % 2 == 0:
            shade_cell(cell, 'F8F9FA')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run(f'Total dosen terverifikasi: {len(all_lecs)} dosen (15 Prodi SI + 10 Prodi Bisnis Digital)')
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0, 97, 0)

doc.add_page_break()

# --- L1.4 Rekap Basis Data Terstruktur ---
subheading(doc, 'Lampiran 1.4 — Rekap Basis Data Dosen Terstruktur (lecturers.json)')
normal(doc, 'Data tersimpan dalam format JSON terstruktur per program studi, siap dikonsumsi oleh sistem SITRIA.', italic=True)
doc.add_paragraph()

# Summary table
tbl_sum = doc.add_table(rows=1, cols=4)
tbl_sum.style = 'Table Grid'
tbl_header(tbl_sum, ['Program Studi', 'Jumlah Dosen', 'Format Data', 'Lokasi Penyimpanan'])
add_row(tbl_sum, ['Sistem Informasi', '15 dosen', 'JSON Terstruktur', 'src/data/prodi/sistem-informasi/'],
        aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])
add_row(tbl_sum, ['Bisnis Digital', '10 dosen', 'JSON Terstruktur', 'src/data/prodi/bisnis-digital/'],
        aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
        bg='EBF5FB')
add_row(tbl_sum, ['TOTAL', '25 dosen aktif', '-', '-'],
        aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
        bold_cols=[0, 1], bg='D5E8D4')

doc.add_paragraph()
normal(doc, 'Setiap entri data dosen memuat: Nama Lengkap, SINTA ID, URL Profil SINTA, dan Google Scholar ID.')

doc.add_page_break()

# --- L1.5 Lembar Validasi Kaprodi ---
subheading(doc, 'Lampiran 1.5 — Lembar Validasi Data Dosen oleh Koordinator Program Studi')
doc.add_paragraph()

tbl_val = doc.add_table(rows=1, cols=1)
tbl_val.style = 'Table Grid'
cell = tbl_val.rows[0].cells[0]
shade_cell(cell, 'FDFEFE')

val_lines = [
    ('LEMBAR VALIDASI DATA DOSEN', True, 12, (0, 51, 102)),
    ('Sistem Informasi Tridharma Akademik (SITRIA)', True, 10, (0, 51, 102)),
    ('', False, 9, (0,0,0)),
    ('Yang bertanda tangan di bawah ini:', False, 10, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('Nama              : ___________________________', False, 10, (50,50,50)),
    ('Jabatan           : Koordinator Program Studi ___________________________', False, 10, (50,50,50)),
    ('Unit Kerja        : FSTI - Institut Teknologi Kalimantan', False, 10, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('Menyatakan bahwa data dosen aktif Program Studi ___________________________ yang telah', False, 10, (50,50,50)),
    ('dikumpulkan dan diverifikasi oleh Saudara Aidil Saputra Kirsan untuk keperluan pengembangan', False, 10, (50,50,50)),
    ('sistem SITRIA adalah BENAR dan SESUAI dengan data resmi program studi.', False, 10, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('Rincian data yang divalidasi:', True, 10, (0, 51, 102)),
    ('  [ ] Nama lengkap seluruh dosen aktif sudah benar', False, 10, (50,50,50)),
    ('  [ ] SINTA ID setiap dosen sudah terverifikasi', False, 10, (50,50,50)),
    ('  [ ] Jumlah dosen sudah sesuai dengan data kepegawaian prodi', False, 10, (50,50,50)),
    ('  [ ] Tidak ada dosen aktif yang terlewat', False, 10, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('Balikpapan, ___________________ 2026', False, 10, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('Koordinator Program Studi,', False, 10, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('', False, 9, (0,0,0)),
    ('', False, 9, (0,0,0)),
    ('(___________________________)', False, 10, (50,50,50)),
    ('NIP. ___________________________', False, 10, (50,50,50)),
]
first = True
for text, bold, size, color in val_lines:
    if first:
        p = cell.paragraphs[0]
        first = False
    else:
        p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold and size >= 11 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = RGBColor(*color)

doc.add_paragraph()
normal(doc, '* Buat 2 lembar: satu untuk Kaprodi SI, satu untuk Kaprodi BD.', italic=True, size=9)

doc.add_page_break()


# ============================================================
# LAMPIRAN KEGIATAN 2
# ============================================================
heading(doc, 'LAMPIRAN KEGIATAN 2')
normal(doc, 'Pengembangan Sistem Pengambilan Data SINTA Otomatis | 14 - 27 Maret 2026', italic=True)
doc.add_paragraph()

# --- L2.1 Dokumen Analisis Struktur SINTA ---
subheading(doc, 'Lampiran 2.1 — Dokumen Analisis Struktur Data Portal SINTA')
doc.add_paragraph()

tbl_a = doc.add_table(rows=1, cols=1)
tbl_a.style = 'Table Grid'
cell = tbl_a.rows[0].cells[0]
shade_cell(cell, 'FDFEFE')

analisis_lines = [
    ('DOKUMEN ANALISIS STRUKTUR DATA PORTAL SINTA', True, 11, (0, 51, 102)),
    ('Kementerian Pendidikan, Kebudayaan, Riset, dan Teknologi', False, 9, (89, 89, 89)),
    ('', False, 9, (0,0,0)),
    ('A. Informasi Portal', True, 10, (0, 51, 102)),
    ('  URL         : https://sinta.kemdiktisaintek.go.id', False, 9, (50,50,50)),
    ('  Tipe Akses  : Publik (tanpa autentikasi untuk data profil dosen)', False, 9, (50,50,50)),
    ('  Data Target : Profil dosen — statistik publikasi, penelitian, pengabdian', False, 9, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('B. Struktur Halaman Profil Dosen yang Dianalisis', True, 10, (0, 51, 102)),
    ('  URL Profil  : sinta.kemdiktisaintek.go.id/authors/profile/{SINTA_ID}', False, 9, (50,50,50)),
    ('  Elemen Data :', False, 9, (50,50,50)),
    ('    - Nama dosen, institusi, program studi', False, 9, (50,50,50)),
    ('    - Skor SINTA, H-Index (Google Scholar, Scopus, WoS)', False, 9, (50,50,50)),
    ('    - Jumlah artikel: Scopus, Google Scholar, Garuda, WoS', False, 9, (50,50,50)),
    ('    - Data penelitian: judul, tahun, sumber dana, jumlah dana', False, 9, (50,50,50)),
    ('    - Data pengabdian: judul, tahun, sumber dana, jumlah dana', False, 9, (50,50,50)),
    ('    - HKI, Buku, Prosiding', False, 9, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('C. Metode Pengambilan Data', True, 10, (0, 51, 102)),
    ('  Metode      : HTTP Request ke endpoint publik + parsing HTML (BeautifulSoup4)', False, 9, (50,50,50)),
    ('  Bahasa      : Python 3.x', False, 9, (50,50,50)),
    ('  Library     : requests, BeautifulSoup4, json, time (rate limiting)', False, 9, (50,50,50)),
    ('  Output      : File JSON terstruktur per program studi', False, 9, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('D. Hasil Analisis & Keputusan Teknis', True, 10, (0, 51, 102)),
    ('  - Seluruh data yang dibutuhkan tersedia secara publik tanpa login', False, 9, (50,50,50)),
    ('  - Jeda antar-request (rate limiting) diterapkan untuk menghormati server SINTA', False, 9, (50,50,50)),
    ('  - Data di-scrape per SINTA ID yang telah diverifikasi pada Kegiatan 1', False, 9, (50,50,50)),
    ('  - Output JSON mengikuti skema standar yang kompatibel dengan frontend Vue.js', False, 9, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('Tanggal Analisis : ___________________ Maret 2026', False, 9, (50,50,50)),
    ('Disusun oleh     : Aidil Saputra Kirsan', False, 9, (50,50,50)),
]
first = True
for text, bold, size, color in analisis_lines:
    if first:
        p = cell.paragraphs[0]
        first = False
    else:
        p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold and size >= 11 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = RGBColor(*color)

doc.add_paragraph()
doc.add_page_break()

# --- L2.2 Rekap Hasil Pipeline Scraping ---
subheading(doc, 'Lampiran 2.2 — Rekap Hasil Data Scraping SINTA (Kedua Prodi)')
normal(doc, 'Tabel berikut menunjukkan data yang berhasil dikumpulkan oleh sistem scraper untuk setiap dosen.', italic=True)
doc.add_paragraph()

subheading(doc, 'A. Rekapitulasi Data Prodi Sistem Informasi', color=(26, 82, 118))
tbl_r1 = doc.add_table(rows=1, cols=7)
tbl_r1.style = 'Table Grid'
tbl_header(tbl_r1, ['No', 'Nama Dosen', 'Scopus', 'Google Scholar', 'H-Index', 'Penelitian', 'PKM'])
for i, (name, sid, sc, gs, h, res, pkm) in enumerate(SI_LECTURERS, 1):
    add_row(tbl_r1, [i, name, sc, gs, h, res, pkm],
            aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER]*5,
            bg='F8F9FA' if i % 2 == 0 else None)

# Total row SI
si_totals = [sum(x[j] for x in SI_LECTURERS) for j in range(2, 7)]
row = tbl_r1.add_row()
vals = ['', 'TOTAL', si_totals[0], si_totals[1], '-', si_totals[3], si_totals[4]]
for j, (cell, val) in enumerate(zip(row.cells, vals)):
    cell.text = str(val)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = cell.paragraphs[0].runs[0]
    run.bold = True; run.font.size = Pt(9)
    shade_cell(cell, 'D5E8D4')

doc.add_paragraph()
subheading(doc, 'B. Rekapitulasi Data Prodi Bisnis Digital', color=(26, 82, 118))
tbl_r2 = doc.add_table(rows=1, cols=7)
tbl_r2.style = 'Table Grid'
tbl_header(tbl_r2, ['No', 'Nama Dosen', 'Scopus', 'Google Scholar', 'H-Index', 'Penelitian', 'PKM'])
for i, (name, sid, sc, gs, h, res, pkm) in enumerate(BD_LECTURERS, 1):
    add_row(tbl_r2, [i, name, sc, gs, h, res, pkm],
            aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER]*5,
            bg='EBF5FB' if i % 2 == 0 else None)

bd_totals = [sum(x[j] for x in BD_LECTURERS) for j in range(2, 7)]
row = tbl_r2.add_row()
vals = ['', 'TOTAL', bd_totals[0], bd_totals[1], '-', bd_totals[3], bd_totals[4]]
for j, (cell, val) in enumerate(zip(row.cells, vals)):
    cell.text = str(val)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = cell.paragraphs[0].runs[0]
    run.bold = True; run.font.size = Pt(9)
    shade_cell(cell, 'DAE8FC')

doc.add_paragraph()

# Summary gabungan
subheading(doc, 'C. Ringkasan Gabungan 2 Prodi', color=(26, 82, 118))
tbl_g = doc.add_table(rows=1, cols=3)
tbl_g.style = 'Table Grid'
tbl_header(tbl_g, ['Indikator', 'Prodi SI', 'Prodi BD'])
gabungan = [
    ('Jumlah Dosen', '15 dosen', '10 dosen'),
    ('Total Penelitian', '86 judul', '46 judul'),
    ('Total PKM', '62 kegiatan', '34 kegiatan'),
    ('Total Dana Penelitian', 'Rp 1.688.119.400', 'Rp 941.172.000'),
    ('Total Dana PKM', 'Rp 599.417.345', 'Rp 602.702.000'),
    ('Rata-rata Penelitian/Dosen', '5,73', '4,60'),
]
for i, (ind, si, bd) in enumerate(gabungan):
    row = tbl_g.add_row()
    for j, (cell, val) in enumerate(zip(row.cells, [ind, si, bd])):
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.bold = (j == 0)
        if i % 2 == 0:
            shade_cell(cell, 'F8F9FA')

doc.add_paragraph()
doc.add_page_break()

# Screenshots K2
screenshot_box(doc, 'bukti_k2_01_terminal_scraper.png',
               'Screenshot terminal saat menjalankan perintah: python sinta_scraper.py (tampilkan output prosesnya)',
               2, '3')
screenshot_box(doc, 'bukti_k2_02_output_json.png',
               'Screenshot isi file sinta_data.json yang terbuka di VSCode/text editor (tampilkan struktur JSON-nya)',
               2, '4')
screenshot_box(doc, 'bukti_k2_03_readme_pipeline.png',
               'Screenshot file README.md yang terbuka di browser/VSCode (tampilkan bagian "Fitur Utama" dan "Alur Kerja")',
               2, '5')

doc.add_page_break()


# ============================================================
# LAMPIRAN KEGIATAN 3
# ============================================================
heading(doc, 'LAMPIRAN KEGIATAN 3')
normal(doc, 'Implementasi Fitur Analitik & Dashboard SITRIA | 21 Maret - 10 April 2026', italic=True)
doc.add_paragraph()

normal(doc,
    'Bukti kegiatan 3 berupa tangkapan layar (screenshot) dari setiap fitur yang telah diimplementasikan. '
    'Jalankan aplikasi terlebih dahulu dengan perintah: npm run dev (di folder presentasi-laboratorium-inovasi-digital), '
    'kemudian buka browser ke alamat localhost:5173.',
    italic=True, size=9)

doc.add_paragraph()

screenshots_k3 = [
    ('bukti_k3_01_dashboard_utama.png',
     'Screenshot halaman Dashboard Utama (localhost:5173/dashboard). Pilih salah satu prodi terlebih dahulu. '
     'Tampilkan statistik lengkap: jumlah publikasi, penelitian, pengabdian, HKI, dan grafik produktivitas.', '1'),
    ('bukti_k3_02_ai_clustering.png',
     'Screenshot halaman Kolaborasi Riset / AI Clustering (localhost:5173/collaboration). '
     'Pastikan kartu-kartu kluster terlihat dengan badge High/Medium/Low kolaborasi.', '2'),
    ('bukti_k3_03_sankey_timeline.png',
     'Screenshot halaman Roadmap Riset / Sankey Timeline (localhost:5173/roadmap). '
     'Pastikan diagram Sankey terlihat dengan alur topik riset dari tahun ke tahun.', '3'),
    ('bukti_k3_04_research_gallery.png',
     'Screenshot halaman Galeri Karya (localhost:5173/gallery). '
     'Tampilkan tab kategori (Penelitian/Pengabdian/Scopus/SINTA/Buku/HKI) dan beberapa karya.', '4'),
    ('bukti_k3_05_expertise_finder.png',
     'Screenshot halaman Expertise Finder (localhost:5173/expertise). '
     'Masukkan contoh topik riset di kolom pencarian dan tampilkan hasil pencocokan pakar.', '5'),
]

for filename, caption, tahap in screenshots_k3:
    screenshot_box(doc, filename, caption, 3, tahap)

doc.add_page_break()


# ============================================================
# LAMPIRAN KEGIATAN 4
# ============================================================
heading(doc, 'LAMPIRAN KEGIATAN 4')
normal(doc, 'Implementasi Dashboard Akreditasi - DTPS & Pendanaan Riset | 28 Maret - 17 April 2026', italic=True)
doc.add_paragraph()

# --- L4.1 Rekap Dana Riset (dari prodiSummary) ---
subheading(doc, 'Lampiran 4.1 — Rekap Data Pendanaan Riset & PKM (Dasar Kalkulasi DTPS)')
normal(doc, 'Data pendanaan berikut bersumber dari sistem SITRIA yang telah berhasil mengumpulkan data dari portal SINTA.', italic=True)
doc.add_paragraph()

tbl_dtps = doc.add_table(rows=1, cols=5)
tbl_dtps.style = 'Table Grid'
tbl_header(tbl_dtps, ['Indikator LKPS', 'Prodi SI', 'Prodi BD', 'Gabungan', 'Satuan'])

dtps_rows = [
    ('Jumlah DTPS', '15', '10', '25', 'Dosen'),
    ('Total Penelitian', '86', '46', '132', 'Judul'),
    ('Rasio Penelitian/DTPS', '5,73', '4,60', '5,28', 'Judul/Dosen'),
    ('Total PKM', '62', '34', '96', 'Kegiatan'),
    ('Rasio PKM/DTPS', '4,13', '3,40', '3,84', 'Kegiatan/Dosen'),
    ('Total Dana Penelitian', 'Rp 1.688.119.400', 'Rp 941.172.000', 'Rp 2.629.291.400', 'Rupiah'),
    ('Dana Penelitian/DTPS', 'Rp 112.541.293', 'Rp 94.117.200', 'Rp 105.171.656', 'Rupiah/Dosen'),
    ('Total Dana PKM', 'Rp 599.417.345', 'Rp 602.702.000', 'Rp 1.202.119.345', 'Rupiah'),
]
for i, row_data in enumerate(dtps_rows):
    row = tbl_dtps.add_row()
    aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
              WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
    for j, (cell, val, align) in enumerate(zip(row.cells, row_data, aligns)):
        cell.text = val
        cell.paragraphs[0].alignment = align
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        run.bold = (j == 0)
        if i % 2 == 0:
            shade_cell(cell, 'F8F9FA')
        if j == 3:
            run.font.color.rgb = RGBColor(0, 51, 102)
            run.bold = True

doc.add_paragraph()

# Screenshots K4
screenshot_box(doc, 'bukti_k4_01_dtps_dashboard.png',
               'Screenshot halaman Dashboard Akreditasi DTPS (localhost:5173/accreditation). '
               'Tampilkan tabel rasio DTPS dengan kalkulasi otomatis dan fitur simulasi dosen lintas prodi.',
               4, '2')
screenshot_box(doc, 'bukti_k4_02_funding_dashboard.png',
               'Screenshot halaman Dashboard Dana & Hibah (localhost:5173/funding). '
               'Tampilkan grafik tren pendanaan dan distribusi sumber dana (internal vs BIMA).',
               4, '3')

doc.add_paragraph()
doc.add_page_break()

# --- L4.5 Template Berita Acara Validasi WD ---
subheading(doc, 'Lampiran 4.5 — Template Berita Acara Validasi DTPS (Menunggu Pelaksanaan)')
normal(doc, 'Dokumen ini akan ditandatangani setelah validasi bersama Wakil Dekan Akademik dilaksanakan (target: sebelum 17 April 2026).', italic=True)
doc.add_paragraph()

tbl_ba = doc.add_table(rows=1, cols=1)
tbl_ba.style = 'Table Grid'
cell = tbl_ba.rows[0].cells[0]
shade_cell(cell, 'FFF9F0')

ba_lines = [
    ('BERITA ACARA VALIDASI DATA DTPS', True, 12, (0, 51, 102)),
    ('Dashboard Akreditasi SITRIA - Sistem Informasi Tridharma Akademik', True, 10, (0, 51, 102)),
    ('', False, 9, (0,0,0)),
    ('Pada hari ini, ___________________________, tanggal ___________________ April 2026,', False, 10, (50,50,50)),
    ('telah dilakukan validasi data Dosen Tetap Program Studi (DTPS) yang ditampilkan pada', False, 10, (50,50,50)),
    ('Dashboard Akreditasi sistem SITRIA.', False, 10, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('Validator:', True, 10, (0, 51, 102)),
    ('  Nama     : ___________________________', False, 10, (50,50,50)),
    ('  Jabatan  : Wakil Dekan Bidang Akademik FSTI - ITK', False, 10, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('Hasil Validasi:', True, 10, (0, 51, 102)),
    ('  [ ] Data jumlah DTPS per program studi sudah sesuai', False, 10, (50,50,50)),
    ('  [ ] Kalkulasi Rasio Penelitian/DTPS sudah sesuai standar LKPS BAN-PT', False, 10, (50,50,50)),
    ('  [ ] Kalkulasi Rasio PKM/DTPS sudah sesuai standar LKPS BAN-PT', False, 10, (50,50,50)),
    ('  [ ] Data dana penelitian dan PKM sudah sesuai dengan rekap resmi', False, 10, (50,50,50)),
    ('  [ ] Dashboard dapat digunakan sebagai alat bantu persiapan akreditasi', False, 10, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('Catatan/Masukan: _______________________________________________', False, 10, (50,50,50)),
    ('________________________________________________________________', False, 10, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('Balikpapan, ___________________ 2026', False, 10, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('Wakil Dekan Akademik FSTI,                    Pembuat Sistem,', False, 10, (50,50,50)),
    ('', False, 9, (0,0,0)),
    ('', False, 9, (0,0,0)),
    ('', False, 9, (0,0,0)),
    ('(________________________)                    (Aidil Saputra Kirsan)', False, 10, (50,50,50)),
    ('NIP. ________________________                 NIP. ________________________', False, 10, (50,50,50)),
]
first = True
for text, bold, size, color in ba_lines:
    if first:
        p = cell.paragraphs[0]
        first = False
    else:
        p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold and size >= 11 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = RGBColor(*color)

doc.add_paragraph()
doc.add_page_break()

# ============================================================
# DAFTAR SCREENSHOT YANG DIBUTUHKAN
# ============================================================
heading(doc, 'DAFTAR SCREENSHOT / FOTO YANG PERLU DILENGKAPI')
normal(doc, 'Simpan semua file gambar ke dalam folder: CPNS/laporan-progress/foto-bukti/', italic=True)
doc.add_paragraph()

tbl_ss = doc.add_table(rows=1, cols=4)
tbl_ss.style = 'Table Grid'
tbl_header(tbl_ss, ['Nama File', 'Kegiatan', 'Cara Mengambil', 'Status'])

screenshots_list = [
    ('bukti_k1_01_foto_koordinasi.jpg', 'K1 - T1', 'Foto fisik saat rapat koordinasi dengan Kaprodi SI dan BD', 'Perlu Foto'),
    ('bukti_k1_02_sinta_profil_dosen.png', 'K1 - T3', 'Screenshot profil salah satu dosen di sinta.kemdiktisaintek.go.id', 'Screenshot'),
    ('bukti_k2_01_terminal_scraper.png', 'K2 - T2', 'Screenshot terminal saat menjalankan: python sinta_scraper.py', 'Screenshot'),
    ('bukti_k2_02_output_json.png', 'K2 - T4', 'Screenshot file sinta_data.json terbuka di VSCode', 'Screenshot'),
    ('bukti_k2_03_readme_pipeline.png', 'K2 - T5', 'Screenshot README.md terbuka di browser/VSCode', 'Screenshot'),
    ('bukti_k3_01_dashboard_utama.png', 'K3 - T4', 'Screenshot localhost:5173/dashboard (pilih prodi dulu)', 'Screenshot'),
    ('bukti_k3_02_ai_clustering.png', 'K3 - T1', 'Screenshot localhost:5173/collaboration (AI Clustering)', 'Screenshot'),
    ('bukti_k3_03_sankey_timeline.png', 'K3 - T2', 'Screenshot localhost:5173/roadmap (Sankey Timeline)', 'Screenshot'),
    ('bukti_k3_04_research_gallery.png', 'K3 - T3', 'Screenshot localhost:5173/gallery (Research Gallery)', 'Screenshot'),
    ('bukti_k3_05_expertise_finder.png', 'K3 - T4', 'Screenshot localhost:5173/expertise (cari topik riset)', 'Screenshot'),
    ('bukti_k4_01_dtps_dashboard.png', 'K4 - T2', 'Screenshot localhost:5173/accreditation (DTPS Dashboard)', 'Screenshot'),
    ('bukti_k4_02_funding_dashboard.png', 'K4 - T3', 'Screenshot localhost:5173/funding (Funding Dashboard)', 'Screenshot'),
]
status_colors = {'Perlu Foto': 'F2DCDB', 'Screenshot': 'EBF5FB'}
for i, (fname, keg, cara, status) in enumerate(screenshots_list):
    row = tbl_ss.add_row()
    vals = [fname, keg, cara, status]
    aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER]
    for j, (cell, val, align) in enumerate(zip(row.cells, vals, aligns)):
        cell.text = val
        cell.paragraphs[0].alignment = align
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        run.bold = (j == 0)
        if j == 3:
            shade_cell(cell, status_colors.get(status, 'FFFFFF'))
            run.font.color.rgb = RGBColor(156, 0, 6) if status == 'Perlu Foto' else RGBColor(26, 82, 118)
            run.bold = True
        elif i % 2 == 0:
            shade_cell(cell, 'F8F9FA')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Catatan: Untuk screenshot aplikasi, jalankan dulu: npm run dev (di folder presentasi-laboratorium-inovasi-digital)')
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(89, 89, 89)

# ============================================================
# SAVE
# ============================================================
output_path = r'D:\Github-ADL\presentasi-umum\CPNS\laporan-progress\Lampiran_Bukti_Aktualisasi_SITRIA.docx'
doc.save(output_path)
print('Saved: ' + output_path)
print(f'\nTotal dosen SI: {len(SI_LECTURERS)}, BD: {len(BD_LECTURERS)}, Gabungan: {len(SI_LECTURERS)+len(BD_LECTURERS)}')
print(f'Total lampiran dihasilkan: Notulensi, Daftar SINTA ID, Verifikasi, Rekap Data, Lembar Validasi, Analisis SINTA, Rekap Scraping, Template BA WD')
print(f'Total screenshot dibutuhkan: {len(screenshots_list)}')
