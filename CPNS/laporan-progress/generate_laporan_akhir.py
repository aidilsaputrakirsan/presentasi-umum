"""
Generator Laporan Akhir Aktualisasi SITRIA — BAGIAN ADMINISTRASI
Pelatihan Dasar CPNS Angkatan VI Tahun 2026
Institut Teknologi Kalimantan

File ini berisi: Halaman Sampul, Lembar Pernyataan Orisinalitas,
Lembar Persetujuan, Lembar Pengesahan, Lembar Konsultasi Mentor & Coach,
OPR (One Page Report), Kata Pengantar, Daftar Isi, Identitas Peserta.

Isi BAB I – BAB III (rancangan, implementasi, penutup) dibuat terpisah
dalam format majalah 2-kolom di `generate_majalah_aktualisasi.py`.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────
#  DESIGN TOKENS
# ─────────────────────────────────────────────

CLR_PRIMARY  = '1B3A6B'   # deep navy
CLR_ACCENT   = '2E86AB'   # steel blue
CLR_GOLD     = 'C9A84C'   # warm gold (cover accent)
CLR_LIGHT    = 'D6EAF8'   # light blue (placeholder header)
CLR_XLIGHT   = 'EBF5FB'   # very light blue
CLR_ALTROW   = 'F4F6F9'   # table alternate row (light gray-blue)
CLR_KGBG     = 'EAF4FB'   # kegiatan title background
CLR_WHITE    = 'FFFFFF'
CLR_COVTOP   = '1B3A6B'   # cover top band
CLR_COVBOT   = '2E86AB'   # cover bottom band

# ─────────────────────────────────────────────
#  HELPER FUNCTIONS
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def shade_para(p, hex_color):
    """Tambahkan background color pada sebuah paragraf."""
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def para_left_border(p, color='2E86AB', sz=24, space=8):
    """Tambahkan left accent bar pada paragraf."""
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left_el = OxmlElement('w:left')
    left_el.set(qn('w:val'), 'single')
    left_el.set(qn('w:sz'), str(sz))
    left_el.set(qn('w:space'), str(space))
    left_el.set(qn('w:color'), color)
    pBdr.append(left_el)
    pPr.append(pBdr)

def stripe_table(table, even_color=CLR_ALTROW):
    """Alternating row colors mulai dari baris data (skip header)."""
    for i, row in enumerate(table.rows[1:], 1):
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_bg(cell, even_color)

def make_cover_band(doc, text, bg_color, text_color='FFFFFF', font_size=11, bold=False,
                    space_before=0, space_after=0, padding_top=6, padding_bot=6):
    """Buat paragraf dengan background color penuh (untuk band di cover)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    # top/bottom cell-like padding via space
    p.paragraph_format.space_before = Pt(padding_top)
    p.paragraph_format.space_after  = Pt(padding_bot)
    shade_para(p, bg_color)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(*bytes.fromhex(text_color))
    return p

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())
    return p

def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    return br

def page_break(doc):
    doc.add_page_break()

def heading(doc, text, level=1, color=None, center=False):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def para(doc, text='', bold=False, italic=False, size=11, align=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
    return p

def add_run(p, text, bold=False, italic=False, size=11):
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return run

def img_placeholder(doc, nomor_gambar, nama_file, keterangan, catatan_ukuran='Screenshot landscape (16:9)'):
    """Placeholder kotak bergaris untuk gambar yang perlu disisipkan manual."""
    table = doc.add_table(rows=3, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Baris header — steel blue tipis
    hdr = table.rows[0].cells[0]
    set_cell_bg(hdr, CLR_ACCENT)
    ph = hdr.paragraphs[0]
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ph.paragraph_format.space_before = Pt(3)
    ph.paragraph_format.space_after  = Pt(3)
    rh = ph.add_run(f'▢  SISIPKAN GAMBAR {nomor_gambar}  ▢')
    rh.bold = True
    rh.font.size = Pt(9)
    rh.font.color.rgb = RGBColor(255, 255, 255)

    # Baris area gambar — sangat terang
    body = table.rows[1].cells[0]
    set_cell_bg(body, CLR_XLIGHT)
    pb = body.paragraphs[0]
    pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pb.paragraph_format.space_before = Pt(28)
    pb.paragraph_format.space_after  = Pt(28)
    r2 = pb.add_run(f'{keterangan}')
    r2.font.size = Pt(10)
    r2.italic = True
    r2.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

    # Baris meta — nama file & catatan ukuran
    meta = table.rows[2].cells[0]
    set_cell_bg(meta, CLR_LIGHT)
    pm = meta.paragraphs[0]
    pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pm.paragraph_format.space_before = Pt(3)
    pm.paragraph_format.space_after  = Pt(3)
    rm = pm.add_run(f'File: {nama_file}   |   {catatan_ukuran}')
    rm.font.size = Pt(8)
    rm.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

    # Caption di bawah
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(3)
    p_cap.paragraph_format.space_after  = Pt(10)
    rc = p_cap.add_run(f'Gambar {nomor_gambar}. {keterangan}')
    rc.italic = True
    rc.font.size = Pt(10)
    rc.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)


def section_title(doc, text, color=CLR_PRIMARY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.45)
    para_left_border(p, color=CLR_ACCENT, sz=24, space=8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_table_row(table, cells_data, bold_col0=True, bg_header=None, font_size=10):
    row = table.add_row()
    for i, cell_data in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = str(cell_data) if cell_data is not None else ''
        for para in cell.paragraphs:
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            for run in para.runs:
                run.font.size = Pt(font_size)
                if i == 0 and bold_col0:
                    run.bold = True
        if bg_header:
            set_cell_bg(cell, bg_header)
    return row

def make_table(doc, headers, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, CLR_PRIMARY)
        for par in cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            par.paragraph_format.space_before = Pt(3)
            par.paragraph_format.space_after  = Pt(3)
            for run in par.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    return table

def add_page_number_footer(doc):
    """Tambahkan nomor halaman di tengah footer setiap section."""
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(0)
        # Garis tipis di atas footer
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        top_f = OxmlElement('w:top')
        top_f.set(qn('w:val'), 'single')
        top_f.set(qn('w:sz'), '4')
        top_f.set(qn('w:space'), '1')
        top_f.set(qn('w:color'), CLR_ACCENT)
        pBdr.append(top_f)
        pPr.append(pBdr)

        def _fld_run(para, instr):
            run = para.add_run()
            run.font.size = Pt(10)
            fc_begin = OxmlElement('w:fldChar')
            fc_begin.set(qn('w:fldCharType'), 'begin')
            instr_el = OxmlElement('w:instrText')
            instr_el.set(qn('xml:space'), 'preserve')
            instr_el.text = f' {instr} '
            fc_end = OxmlElement('w:fldChar')
            fc_end.set(qn('w:fldCharType'), 'end')
            run._r.append(fc_begin)
            run._r.append(instr_el)
            run._r.append(fc_end)
            return run

        r_dash1 = p.add_run('– ')
        r_dash1.font.size = Pt(10)
        _fld_run(p, 'PAGE')
        r_of = p.add_run(' dari ')
        r_of.font.size = Pt(10)
        _fld_run(p, 'NUMPAGES')
        r_dash2 = p.add_run(' –')
        r_dash2.font.size = Pt(10)


def add_toc_entry(doc, text, is_bold=False, indent=0):
    """Entri daftar isi dengan dot leader dan placeholder halaman."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)

    # Tab kanan dengan dot leader di posisi ~13.5 cm dari margin kiri
    pPr = p._p.get_or_add_pPr()
    tabs_el = OxmlElement('w:tabs')
    tab_el = OxmlElement('w:tab')
    tab_el.set(qn('w:val'), 'right')
    tab_el.set(qn('w:leader'), 'dot')
    tab_el.set(qn('w:pos'), '7560')   # 7560 twips ≈ 13.5 cm
    tabs_el.append(tab_el)
    pPr.append(tabs_el)

    run = p.add_run(text)
    run.bold = is_bold
    run.font.size = Pt(11)
    run2 = p.add_run('\t...')
    run2.bold = is_bold
    run2.font.size = Pt(11)
    return p


def horizontal_line(doc):
    """Garis pemisah dua lapis: tebal (navy) + tipis (steel blue)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    # Garis atas tebal — navy
    top_el = OxmlElement('w:top')
    top_el.set(qn('w:val'), 'single')
    top_el.set(qn('w:sz'), '12')
    top_el.set(qn('w:space'), '1')
    top_el.set(qn('w:color'), CLR_PRIMARY)
    pBdr.append(top_el)
    # Garis bawah tipis — steel blue
    bot_el = OxmlElement('w:bottom')
    bot_el.set(qn('w:val'), 'single')
    bot_el.set(qn('w:sz'), '4')
    bot_el.set(qn('w:space'), '1')
    bot_el.set(qn('w:color'), CLR_ACCENT)
    pBdr.append(bot_el)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    return p


# ─────────────────────────────────────────────
#  DATA
# ─────────────────────────────────────────────

DATA = {
    'nama': 'Aidil Saputra Kirsan, S.ST., M.Tr.Kom',
    'nip': '199403172025061004',
    'jabatan': 'Dosen Asisten Ahli / Kepala Lab Inovasi Digital',
    'pangkat': 'Penata Muda Tk.I / III-b',
    'unit_kerja': 'Lab Inovasi Digital, Prodi Sistem Informasi, FSTI ITK',
    'instansi': 'Institut Teknologi Kalimantan',
    'alamat': 'Jl. Soekarno-Hatta KM 15, Balikpapan, Kaltim 76127',
    'mentor_nama': 'Irma Fitria, S.Si., M.Si',
    'mentor_nip': '199303232022032016',
    'mentor_jabatan': 'Wakil Dekan Bidang Akademik dan Kemahasiswaan FSTI ITK',
    'coach_nama': 'Mustari Kurniawati, S.IP., MPA',
    'coach_nip': '197712232005012001',
    'penguji_nama': 'Dr. M. Muhamad Harry Rahmadi, S.Pi., MM.',
    'penguji_nip': '198510092011011012',
    'angkatan': 'VI',
    'kelompok': 'I',
    'tahun': '2026',
    'penyelenggara': 'Pusat Pelatihan dan Pengembangan – LAN RI',
    'judul': 'Pengembangan Sistem Dashboard Analitik Tridharma Berbasis Data Multi-Sumber (SITRIA) pada Lab Inovasi Digital FSTI ITK',
    'periode': '7 Maret – 22 April 2026',
    'tanggal_laporan': '15 April 2026',
    'tanggal_seminar': '21 April 2026',
    'tempat_seminar': 'Online / Daring',
}


# ─────────────────────────────────────────────
#  DOCUMENT SETUP
# ─────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

# Footer: nomor halaman
add_page_number_footer(doc)

# Default paragraph style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)


# ═══════════════════════════════════════════════
#  1. HALAMAN SAMPUL
# ═══════════════════════════════════════════════

# ── Judul ──
p = para(doc, 'LAPORAN AKTUALISASI', bold=True, size=16, align='center', space_before=20, space_after=4)
for run in p.runs: run.font.color.rgb = RGBColor(*bytes.fromhex(CLR_PRIMARY))
p = para(doc, 'PELATIHAN DASAR CPNS ANGKATAN VI TAHUN 2026', bold=True, size=12, align='center', space_after=4)
for run in p.runs: run.font.color.rgb = RGBColor(*bytes.fromhex(CLR_PRIMARY))
p = para(doc, 'PUSAT PELATIHAN DAN PENGEMBANGAN – LAN RI', size=10, align='center', space_after=16)

horizontal_line(doc)

# ── SITRIA title block ──
p = para(doc, 'SITRIA', bold=True, size=26, align='center', space_before=14, space_after=2)
for run in p.runs: run.font.color.rgb = RGBColor(*bytes.fromhex(CLR_PRIMARY))
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after  = Pt(6)
r_sub = p_sub.add_run('Sistem Informasi Tridharma Akademik')
r_sub.bold = True
r_sub.font.size = Pt(14)
r_sub.font.color.rgb = RGBColor(*bytes.fromhex(CLR_ACCENT))
p = para(doc, 'Dashboard Analitik Tridharma Dosen Berbasis Data Multi-Sumber',
         italic=True, size=11, align='center', space_after=2)
p = para(doc, 'Lab Inovasi Digital – Prodi Sistem Informasi & Bisnis Digital',
         italic=True, size=10, align='center', space_after=2)
p = para(doc, 'Fakultas Sains dan Teknologi Industri – Institut Teknologi Kalimantan',
         italic=True, size=10, align='center', space_after=16)

horizontal_line(doc)

# ── Penyusun ──
p = para(doc, 'Disusun oleh:', size=11, align='center', space_before=12, space_after=4)
p = para(doc, DATA['nama'], bold=True, size=13, align='center', space_after=2)
for run in p.runs: run.font.color.rgb = RGBColor(*bytes.fromhex(CLR_PRIMARY))
p = para(doc, f"NIP. {DATA['nip']}", size=11, align='center', space_after=2)
p = para(doc, DATA['jabatan'], size=11, align='center', space_after=2)
p = para(doc, DATA['unit_kerja'], size=11, align='center', space_after=16)

horizontal_line(doc)

p = para(doc, f"Periode Aktualisasi: {DATA['periode']}", size=11, align='center', space_before=10, space_after=2)
p = para(doc, f"Tanggal Laporan: {DATA['tanggal_laporan']}", size=11, align='center', space_after=16)

page_break(doc)


# ═══════════════════════════════════════════════
#  2. LEMBAR PERNYATAAN ORISINALITAS
# ═══════════════════════════════════════════════

heading(doc, 'LEMBAR PERNYATAAN ORISINALITAS LAPORAN AKTUALISASI', level=1, center=True)
horizontal_line(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(10)
add_run(p, 'Saya yang bertanda tangan di bawah ini:', size=12)

# Identitas tabel
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
rows_data = [
    ('Nama', ':', DATA['nama']),
    ('NIP', ':', DATA['nip']),
    ('Jabatan', ':', DATA['jabatan']),
    ('Unit Kerja', ':', DATA['unit_kerja']),
    ('Instansi', ':', DATA['instansi']),
]
for i, (k, sep, v) in enumerate(rows_data):
    row = table.rows[i]
    row.cells[0].text = k
    row.cells[1].text = sep
    row.cells[2].text = v
    row.cells[0].width = Cm(4)
    row.cells[1].width = Cm(0.5)
    for cell in row.cells:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(11)
            par.paragraph_format.space_before = Pt(2)
            par.paragraph_format.space_after = Pt(2)
    row.cells[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(8)
add_run(p, 'Menyatakan bahwa Laporan Aktualisasi yang berjudul ', size=12)
add_run(p, f'"{DATA["judul"]}"', bold=True, size=12)
add_run(p, ' ini adalah hasil karya saya sendiri dan bebas dari plagiarisme. '
           'Kutipan dan referensi yang digunakan telah dicantumkan dengan benar dan sesuai ketentuan yang berlaku. '
           'Apabila terbukti pernyataan ini tidak benar, maka saya bersedia menerima sanksi sesuai peraturan yang berlaku.', size=12)

p = para(doc, f'Balikpapan, {DATA["tanggal_laporan"]}', size=12, align='right', space_before=20, space_after=2)
p = para(doc, 'Yang Menyatakan,', size=12, align='right', space_after=60)
p = para(doc, DATA['nama'], bold=True, size=12, align='right', space_after=2)
p = para(doc, f"NIP. {DATA['nip']}", size=12, align='right')

page_break(doc)


# ═══════════════════════════════════════════════
#  3. LEMBAR PERSETUJUAN
# ═══════════════════════════════════════════════

heading(doc, 'LEMBAR PERSETUJUAN LAPORAN AKTUALISASI', level=1, center=True)
horizontal_line(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_run(p, 'Laporan Aktualisasi ini diajukan oleh:', size=12)

table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
rows_data = [
    ('Nama', ':', DATA['nama']),
    ('NIP', ':', DATA['nip']),
    ('Jabatan', ':', DATA['jabatan']),
    ('Unit Kerja', ':', DATA['unit_kerja']),
    ('Judul', ':', DATA['judul']),
]
for i, (k, sep, v) in enumerate(rows_data):
    row = table.rows[i]
    row.cells[0].text = k
    row.cells[1].text = sep
    row.cells[2].text = v
    row.cells[0].width = Cm(4)
    row.cells[1].width = Cm(0.5)
    for cell in row.cells:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(11)
            par.paragraph_format.space_before = Pt(2)
            par.paragraph_format.space_after = Pt(2)
    row.cells[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_run(p, 'Telah mendapat persetujuan untuk diseminarkan sebagai Laporan Aktualisasi '
           f'Pelatihan Dasar CPNS Angkatan {DATA["angkatan"]} Tahun {DATA["tahun"]}.', size=12)

p = para(doc, f'Balikpapan, {DATA["tanggal_laporan"]}', size=12, space_before=20, space_after=4)

# TTD grid
table_ttd = doc.add_table(rows=1, cols=2)
table_ttd.alignment = WD_TABLE_ALIGNMENT.CENTER
cells_ttd = table_ttd.rows[0].cells
for i, (role, nama, nip) in enumerate([
    ('Menyetujui, Coach', DATA['coach_nama'], f"NIP. {DATA['coach_nip']}"),
    ('Mengetahui, Mentor', DATA['mentor_nama'], f"NIP. {DATA['mentor_nip']}"),
]):
    c = cells_ttd[i]
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(c.paragraphs[0], role, size=11)
    p2 = c.add_paragraph('\n\n\n')
    p3 = c.add_paragraph(nama)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.runs[0].bold = True
    p3.runs[0].font.size = Pt(11)
    p4 = c.add_paragraph(nip)
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.runs[0].font.size = Pt(11)

page_break(doc)


# ═══════════════════════════════════════════════
#  4. LEMBAR PENGESAHAN
# ═══════════════════════════════════════════════

heading(doc, 'LEMBAR PENGESAHAN LAPORAN AKTUALISASI', level=1, center=True)
horizontal_line(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_run(p, 'Laporan Aktualisasi ini telah diseminarkan di hadapan penguji dan dinyatakan '
           'telah memenuhi syarat untuk diterima.', size=12)

table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
rows_data = [
    ('Nama', ':', DATA['nama']),
    ('NIP', ':', DATA['nip']),
    ('Jabatan', ':', DATA['jabatan']),
    ('Unit Kerja', ':', DATA['unit_kerja']),
    ('Tanggal Seminar', ':', DATA['tanggal_seminar']),
    ('Tempat', ':', DATA['tempat_seminar']),
]
for i, (k, sep, v) in enumerate(rows_data):
    row = table.rows[i]
    row.cells[0].text = k
    row.cells[1].text = sep
    row.cells[2].text = v
    row.cells[0].width = Cm(4.5)
    row.cells[1].width = Cm(0.5)
    for cell in row.cells:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(11)
            par.paragraph_format.space_before = Pt(2)
            par.paragraph_format.space_after = Pt(2)
    row.cells[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

table_ttd = doc.add_table(rows=1, cols=3)
table_ttd.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (role, nama, nip) in enumerate([
    ('Penguji', DATA['penguji_nama'], f"NIP. {DATA['penguji_nip']}"),
    ('Coach', DATA['coach_nama'], f"NIP. {DATA['coach_nip']}"),
    ('Mentor', DATA['mentor_nama'], f"NIP. {DATA['mentor_nip']}"),
]):
    c = table_ttd.rows[0].cells[i]
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(c.paragraphs[0], role, bold=True, size=11)
    p2 = c.add_paragraph('\n\n\n')
    p3 = c.add_paragraph(nama)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.runs[0].bold = True
    p3.runs[0].font.size = Pt(10)
    p4 = c.add_paragraph(nip)
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.runs[0].font.size = Pt(10)

page_break(doc)


# ═══════════════════════════════════════════════
#  5. LEMBAR KONSULTASI MENTOR
# ═══════════════════════════════════════════════

heading(doc, 'LEMBAR KONSULTASI MENTOR', level=1, center=True)
horizontal_line(doc)

p = doc.add_paragraph()
add_run(p, 'Nama Mentor\t: ', bold=True, size=11)
add_run(p, DATA['mentor_nama'], size=11)
p2 = doc.add_paragraph()
add_run(p2, 'NIP\t\t\t: ', bold=True, size=11)
add_run(p2, DATA['mentor_nip'], size=11)
p3 = doc.add_paragraph()
add_run(p3, 'Jabatan\t\t: ', bold=True, size=11)
add_run(p3, DATA['mentor_jabatan'], size=11)
doc.add_paragraph()

tbl = make_table(doc,
    ['No', 'Tanggal', 'Uraian Konsultasi / Catatan', 'Paraf Mentor'],
    col_widths=[1, 3, 9, 3])

konsultasi_mentor = [
    ('1', '____________ 2026', 'Konsultasi persetujuan rancangan aktualisasi dan dukungan pimpinan unit kerja', ''),
    ('2', '____________ 2026', 'Penyelarasan gagasan "SITRIA" dan batasan sasaran dashboard', ''),
    ('3', '____________ 2026', 'Review progres Kegiatan 1 dan 2: data dosen dan sistem pengambilan data SINTA', ''),
    ('4', '____________ 2026', 'Review progres Kegiatan 3 dan 4: fitur analitik dan dashboard akreditasi', ''),
    ('5', '____________ 2026', 'Konsultasi final laporan aktualisasi dan persiapan seminar', ''),
]
for row_data in konsultasi_mentor:
    r = tbl.add_row()
    for i, val in enumerate(row_data):
        r.cells[i].text = val
        r.cells[i].paragraphs[0].paragraph_format.space_before = Pt(2)
        r.cells[i].paragraphs[0].paragraph_format.space_after = Pt(2)
        for run in r.cells[i].paragraphs[0].runs:
            run.font.size = Pt(10)
        if i == 0:
            r.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

page_break(doc)


# ═══════════════════════════════════════════════
#  6. LEMBAR KONSULTASI COACH
# ═══════════════════════════════════════════════

heading(doc, 'LEMBAR KONSULTASI COACH', level=1, center=True)
horizontal_line(doc)

p = doc.add_paragraph()
add_run(p, 'Nama Coach\t: ', bold=True, size=11)
add_run(p, DATA['coach_nama'], size=11)
p2 = doc.add_paragraph()
add_run(p2, 'NIP\t\t\t: ', bold=True, size=11)
add_run(p2, DATA['coach_nip'], size=11)
doc.add_paragraph()

tbl = make_table(doc,
    ['No', 'Tanggal', 'Uraian Konsultasi / Catatan', 'Paraf Coach'],
    col_widths=[1, 3, 9, 3])

konsultasi_coach = [
    ('1', '____________ 2026', 'Konsultasi awal: pemaparan rancangan dan gagasan kreatif SITRIA', ''),
    ('2', '____________ 2026', 'Review usulan kegiatan kritis pemecahan masalah (gagasan kreatif "SITRIA")', ''),
    ('3', '____________ 2026', 'Penguatan relevansi nilai BerAKHLAK pada setiap kegiatan aktualisasi', ''),
    ('4', '____________ 2026', 'Review implementasi Kegiatan 3–5 dan bukti-bukti pendukung', ''),
    ('5', '____________ 2026', 'Konsultasi penyusunan laporan akhir dan persiapan seminar aktualisasi', ''),
]
for row_data in konsultasi_coach:
    r = tbl.add_row()
    for i, val in enumerate(row_data):
        r.cells[i].text = val
        r.cells[i].paragraphs[0].paragraph_format.space_before = Pt(2)
        r.cells[i].paragraphs[0].paragraph_format.space_after = Pt(2)
        for run in r.cells[i].paragraphs[0].runs:
            run.font.size = Pt(10)
        if i == 0:
            r.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

page_break(doc)


# ═══════════════════════════════════════════════
#  7. OPR — ONE PAGE REPORT
# ═══════════════════════════════════════════════

heading(doc, 'OPR — ONE PAGE REPORT', level=1, center=True)
horizontal_line(doc)

opr_table = doc.add_table(rows=0, cols=2)
opr_table.style = 'Table Grid'
opr_table.alignment = WD_TABLE_ALIGNMENT.CENTER

def opr_row(table, label, value):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[0].width = Cm(5.5)
    row.cells[1].text = value
    for par in row.cells[0].paragraphs:
        par.runs[0].bold = True
        par.runs[0].font.size = Pt(10)
        par.paragraph_format.space_before = Pt(2)
        par.paragraph_format.space_after = Pt(2)
    for par in row.cells[1].paragraphs:
        for run in par.runs:
            run.font.size = Pt(10)
        par.paragraph_format.space_before = Pt(2)
        par.paragraph_format.space_after = Pt(2)
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

opr_row(opr_table, 'Nama Peserta', DATA['nama'])
opr_row(opr_table, 'NIP', DATA['nip'])
opr_row(opr_table, 'Jabatan', DATA['jabatan'])
opr_row(opr_table, 'Unit Kerja', DATA['unit_kerja'])
opr_row(opr_table, 'Instansi', DATA['instansi'])
opr_row(opr_table, 'Angkatan/Kelompok', f"Angkatan {DATA['angkatan']} / Kelompok {DATA['kelompok']}")
opr_row(opr_table, 'Mentor', DATA['mentor_nama'])
opr_row(opr_table, 'Coach', DATA['coach_nama'])
opr_row(opr_table, 'Judul Aktualisasi', DATA['judul'])
opr_row(opr_table, 'Isu Terpilih',
        'Pengelolaan Data Riset dan Pengabdian Dosen di Lab Inovasi Digital Belum Terpusat dan Teranalisis '
        '(Skor USG: 14 — Urgency=5, Seriousness=5, Growth=4)')
opr_row(opr_table, 'Gagasan Kreatif',
        'SITRIA (Sistem Informasi Tridharma Akademik): Dashboard analitik berbasis web yang mengintegrasikan data '
        'riset, publikasi, dan pengabdian dosen secara otomatis dari portal SINTA Kemdiktisaintek, '
        'dilengkapi kemampuan kecerdasan buatan untuk pengelompokan dan analisis topik riset secara otomatis.')
opr_row(opr_table, 'Kegiatan Utama (5)',
        '1. Pemetaan & Pengumpulan Data Dosen (07–20 Mar)\n'
        '2. Pengembangan Sistem Pengambilan Data SINTA Otomatis (14–27 Mar)\n'
        '3. Implementasi Fitur Analitik & Dashboard (21 Mar – 10 Apr)\n'
        '4. Dashboard Akreditasi DTPS & Pendanaan Riset (28 Mar – 17 Apr)\n'
        '5. Peluncuran, Sosialisasi & Evaluasi SITRIA (11–22 Apr)')
opr_row(opr_table, 'Output Utama',
        '• 25 dosen terdata dari 2 prodi (SI & Bisnis Digital)\n'
        '• 7 fitur dashboard aktif: Analitik Utama, Research Gallery, AI Clustering, Sankey Timeline, '
        'Funding Dashboard, DTPS Akreditasi, Expertise Finder\n'
        '• Sistem diluncurkan ke server Lab Inovasi Digital, 15 April 2026\n'
        '• 23 peserta sosialisasi dari kalangan dosen, Kaprodi, dan laboran')
opr_row(opr_table, 'Nilai BerAKHLAK',
        'Berorientasi Pelayanan (data real-time bagi pimpinan), Akuntabel (kalkulasi DTPS otomatis), '
        'Kompeten (pengembangan sistem informasi & kecerdasan buatan), Harmonis (kolaborasi lintas prodi), Loyal (mendukung akreditasi ITK), '
        'Adaptif (data multi-sumber, AI/ML), Kolaboratif (melibatkan semua pemangku kepentingan)')
opr_row(opr_table, 'Manfaat Utama',
        '• Bagi Prodi SI & Bisnis Digital: persiapan borang akreditasi lebih cepat dan akurat\n'
        '• Bagi pimpinan FSTI: monitoring kinerja Tridharma dosen secara real-time\n'
        '• Bagi dosen: visibilitas profil riset dan deteksi peluang kolaborasi\n'
        '• Bagi institusi: model best practice yang berpotensi direplikasi ke prodi lain')
opr_row(opr_table, 'Status Penyelesaian',
        'SELESAI 100% — Semua 5 kegiatan dan 25 tahapan aktualisasi berhasil diselesaikan dalam periode 7 Maret – 22 April 2026')

page_break(doc)


# ═══════════════════════════════════════════════
#  8. KATA PENGANTAR
# ═══════════════════════════════════════════════

heading(doc, 'KATA PENGANTAR', level=1, center=True)
horizontal_line(doc)

kata_pengantar_paras = [
    ('Puji syukur penulis panjatkan kepada Tuhan Yang Maha Esa atas limpahan rahmat dan karunia-Nya '
     'sehingga penulis dapat menyelesaikan penyusunan Laporan Aktualisasi ini dalam rangka '
     'Pelatihan Dasar CPNS Angkatan VI Tahun 2026 Institut Teknologi Kalimantan.'),
    ('Laporan Aktualisasi berjudul "Pengembangan Sistem Dashboard Analitik Tridharma Berbasis Data '
     'Multi-Sumber (SITRIA) pada Lab Inovasi Digital FSTI Institut Teknologi Kalimantan" ini disusun '
     'sebagai wujud internalisasi dan implementasi nilai-nilai dasar ASN BerAKHLAK serta peran dan '
     'kedudukan PNS dalam mendukung terwujudnya tata kelola pemerintahan yang cerdas.'),
    ('Laporan ini merupakan dokumentasi pelaksanaan seluruh kegiatan aktualisasi selama periode habituasi '
     '7 Maret – 22 April 2026 di Lab Inovasi Digital, Prodi Sistem Informasi, FSTI ITK. Seluruh '
     'rangkaian kegiatan mulai dari pemetaan data dosen, pengembangan sistem pengambilan data otomatis, '
     'implementasi dashboard analitik, hingga peluncuran dan sosialisasi sistem telah berhasil diselesaikan '
     'sesuai rencana yang telah ditetapkan.'),
    ('Penulis mengucapkan terima kasih kepada: (1) Rektor ITK atas kesempatan mengikuti Latsar CPNS; '
     '(2) Ibu Mustari Kurniawati, S.IP., MPA selaku Coach atas bimbingan dan arahan yang diberikan; '
     '(3) Ibu Irma Fitria, S.Si., M.Si selaku Mentor atas dukungan penuh dari lingkungan unit kerja; '
     '(4) Pimpinan FSTI ITK atas dukungan institusional; (5) Koordinator Prodi Sistem Informasi dan '
     'Prodi Bisnis Digital atas kerja sama dalam validasi data dan sosialisasi sistem; '
     '(6) Seluruh rekan peserta Latsar CPNS Angkatan VI Tahun 2026.'),
    ('Semoga Laporan Aktualisasi ini dapat memberikan manfaat nyata bagi peningkatan kinerja dan '
     'kualitas layanan Lab Inovasi Digital, Program Studi Sistem Informasi, dan Program Studi Bisnis '
     'Digital FSTI ITK, serta menjadi kontribusi positif bagi pengembangan tata kelola akademik '
     'Institut Teknologi Kalimantan.'),
]

for text in kata_pengantar_paras:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(1.25)
    add_run(p, text, size=12)

p = para(doc, f'Balikpapan, {DATA["tanggal_laporan"]}', size=12, align='right', space_before=16, space_after=4)
p = para(doc, 'Penulis,', size=12, align='right', space_after=50)
p = para(doc, DATA['nama'], bold=True, size=12, align='right', space_after=2)
p = para(doc, f"NIP. {DATA['nip']}", size=12, align='right')

page_break(doc)


# ═══════════════════════════════════════════════
#  9. DAFTAR ISI
# ═══════════════════════════════════════════════

heading(doc, 'DAFTAR ISI', level=1, center=True)
horizontal_line(doc)

p_note = doc.add_paragraph()
p_note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_note.paragraph_format.space_before = Pt(0)
p_note.paragraph_format.space_after = Pt(6)
rn = p_note.add_run('*) Isi nomor halaman setelah dokumen final dicetak / dibuka di Word')
rn.italic = True
rn.font.size = Pt(9)
rn.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# Entri-entri daftar isi  (teks, is_bold, indent_cm)
# Catatan: BAB I–III disusun dalam MAJALAH AKTUALISASI SITRIA (dokumen terpisah).
daftar_isi = [
    ('Lembar Pernyataan Orisinalitas Laporan Aktualisasi', False, 0),
    ('Lembar Persetujuan Laporan Aktualisasi',             False, 0),
    ('Lembar Pengesahan Laporan Aktualisasi',              False, 0),
    ('Lembar Konsultasi Mentor',                           False, 0),
    ('Lembar Konsultasi Coach',                            False, 0),
    ('OPR (One Page Report)',                              False, 0),
    ('Kata Pengantar',                                     False, 0),
    ('Daftar Isi',                                         False, 0),
    ('Identitas Peserta',                                  False, 0),
    ('Majalah Aktualisasi SITRIA (dokumen terpisah)',      True,  0),
    ('BAB I   Rancangan Aktualisasi',                      False, 0.8),
    ('BAB II  Implementasi Aktualisasi',                   False, 0.8),
    ('BAB III Penutup',                                    False, 0.8),
]

for teks, bold, indent in daftar_isi:
    add_toc_entry(doc, teks, is_bold=bold, indent=indent)

page_break(doc)


# ═══════════════════════════════════════════════
#  10. IDENTITAS PESERTA
# ═══════════════════════════════════════════════

heading(doc, 'IDENTITAS PESERTA', level=1, center=True)
horizontal_line(doc)

id_table = doc.add_table(rows=0, cols=3)
id_table.style = 'Table Grid'
id_rows = [
    ('Nama Lengkap', DATA['nama']),
    ('NIP', DATA['nip']),
    ('Tempat, Tgl Lahir', 'Ujung Pandang, 17 Maret 1994'),
    ('Jenis Kelamin', 'Laki-laki'),
    ('Agama', 'Islam'),
    ('Pendidikan Terakhir', 'Magister Terapan Komputer (M.Tr.Kom)'),
    ('Program Studi', 'Sistem Informasi'),
    ('Jabatan', DATA['jabatan']),
    ('Pangkat/Golongan', DATA['pangkat']),
    ('Unit Kerja', DATA['unit_kerja']),
    ('Instansi', DATA['instansi']),
    ('Alamat Instansi', DATA['alamat']),
    ('Nama Mentor', DATA['mentor_nama']),
    ('Nama Coach', DATA['coach_nama']),
    ('Angkatan/Kelompok', f"Angkatan {DATA['angkatan']} / Kelompok {DATA['kelompok']}"),
]
for label, value in id_rows:
    row = id_table.add_row()
    row.cells[0].text = label
    row.cells[1].text = ':'
    row.cells[2].text = value
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(0.5)
    for cell in row.cells:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(11)
            par.paragraph_format.space_before = Pt(2)
            par.paragraph_format.space_after = Pt(2)
    row.cells[0].paragraphs[0].runs[0].bold = True

page_break(doc)



# ═══════════════════════════════════════════════
#  CATATAN PENTING
# ═══════════════════════════════════════════════
# Isi BAB I – BAB III (Rancangan, Implementasi, Penutup) dibuat dalam
# dokumen terpisah: Majalah Aktualisasi SITRIA.
# Generator: generate_majalah_aktualisasi.py
# Output    : Majalah_Aktualisasi_SITRIA_2026.docx



# ═══════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════

# ═══════════════════════════════════════════════
#  LAMPIRAN GAMBAR — Placeholder per Kegiatan
# ═══════════════════════════════════════════════

# ── Lampiran terpisah dihapus; bukti fisik sudah masuk per kegiatan di BAB II ──

# ═══════════════════════════════════════════════

output_path = r'd:\Github-ADL\presentasi-umum\CPNS\laporan-progress\Laporan_Akhir_Aktualisasi_SITRIA_2026.docx'
doc.save(output_path)
print(f'[OK] Laporan akhir berhasil dibuat: {output_path}')
