"""
Generator Majalah Aktualisasi SITRIA
Pelatihan Dasar CPNS Angkatan VI Tahun 2026
Institut Teknologi Kalimantan

Dokumen ini adalah versi MAJALAH dari BAB I – BAB III laporan aktualisasi.
Format: 2-kolom, tipografi majalah, drop cap, pull-quote, sidebar BerAKHLAK,
cover bab full-bleed. Bagian administrasi (sampul, lembar pengesahan, dst.)
dibuat di `generate_laporan_akhir.py`.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
#  DESIGN TOKENS — MAJALAH
# ─────────────────────────────────────────────

CLR_PRIMARY   = '1B3A6B'   # navy deep — masthead, judul bab
CLR_ACCENT    = 'C0392B'   # merah majalah — pull-quote, aksen
CLR_GOLD      = 'C9A84C'   # gold — ornamen cover
CLR_STEEL     = '2E86AB'   # steel blue — secondary
CLR_CREAM     = 'FAF6EE'   # krem — background sidebar
CLR_PAPER     = 'FFFFFF'
CLR_SOFTGRAY  = 'ECEEF1'   # panel tone
CLR_DARK      = '1A1A1A'   # body text
CLR_MUTE      = '6B6B6B'   # caption / meta
CLR_SIDEBAR   = 'EEF3F8'   # sidebar BerAKHLAK
CLR_SIDEACC   = '1B3A6B'

FONT_BODY   = 'Georgia'       # serif — teks isi (nuansa majalah)
FONT_HEAD   = 'Calibri'       # sans — heading, pull-quote
FONT_DISPL  = 'Calibri'       # display — cover / masthead

# ─────────────────────────────────────────────
#  LOW-LEVEL XML HELPERS
# ─────────────────────────────────────────────

def _oxml(tag, **attrs):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), v)
    return el

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = _oxml('w:shd', **{'w:val': 'clear', 'w:color': 'auto', 'w:fill': hex_color})
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    """Margin dalam cell (twips: 1 cm ≈ 567)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = _oxml('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = _oxml(f'w:{side}', **{'w:w': str(val), 'w:type': 'dxa'})
        tcMar.append(el)
    tcPr.append(tcMar)

def shade_para(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = _oxml('w:shd', **{'w:val': 'clear', 'w:color': 'auto', 'w:fill': hex_color})
    pPr.append(shd)

def para_border(p, top=None, bottom=None, left=None, right=None):
    pPr = p._p.get_or_add_pPr()
    pBdr = _oxml('w:pBdr')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = _oxml(f'w:{side}',
                       **{'w:val': val.get('val', 'single'),
                          'w:sz': str(val.get('sz', 4)),
                          'w:space': str(val.get('space', 4)),
                          'w:color': val.get('color', '000000')})
            pBdr.append(el)
    pPr.append(pBdr)

def set_columns(section, num=2, space_cm=0.7, sep=False):
    """Multi-kolom untuk section (gaya majalah)."""
    sectPr = section._sectPr
    # hapus cols lama jika ada
    for existing in sectPr.findall(qn('w:cols')):
        sectPr.remove(existing)
    cols = _oxml('w:cols',
                 **{'w:num': str(num),
                    'w:space': str(int(space_cm * 567)),
                    'w:equalWidth': '1'})
    if sep:
        cols.set(qn('w:sep'), '1')
    sectPr.append(cols)

def add_section_continuous(doc):
    """Tambah section baru continuous (pindah kolom)."""
    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    return new_section

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = _oxml('w:br', **{'w:type': 'page'})
    run._r.append(br)
    return p

def add_column_break(p):
    run = p.add_run()
    br = _oxml('w:br', **{'w:type': 'column'})
    run._r.append(br)

def insert_hr_line(doc, color=CLR_PRIMARY, sz=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    para_border(p, bottom={'color': color, 'sz': sz, 'space': '1'})
    return p

# ─────────────────────────────────────────────
#  TYPOGRAPHY HELPERS
# ─────────────────────────────────────────────

def _set_font(run, name=FONT_BODY, size=10, bold=False, italic=False,
              color=CLR_DARK, small_caps=False):
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = _oxml('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rFonts.set(qn(attr), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    if small_caps:
        rPr_sc = _oxml('w:smallCaps', **{'w:val': '1'})
        rPr.append(rPr_sc)

def para(doc, text='', font=FONT_BODY, size=10, bold=False, italic=False,
         color=CLR_DARK, align=None, first_indent=0, space_before=0, space_after=4,
         line_spacing=1.25):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_indent:
        pf.first_line_indent = Cm(first_indent)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if text:
        run = p.add_run(text)
        _set_font(run, name=font, size=size, bold=bold, italic=italic, color=color)
    return p

def add_run(p, text, font=FONT_BODY, size=10, bold=False, italic=False, color=CLR_DARK):
    run = p.add_run(text)
    _set_font(run, name=font, size=size, bold=bold, italic=italic, color=color)
    return run

def body_para(doc, text, size=10, align='justify', first_indent=0.6, space_after=4):
    """Paragraf body majalah — Georgia, justify, indent awal."""
    return para(doc, text, font=FONT_BODY, size=size, align=align,
                first_indent=first_indent, space_before=0, space_after=space_after,
                line_spacing=1.3)

# ─────────────────────────────────────────────
#  MAGAZINE COMPONENTS
# ─────────────────────────────────────────────

def drop_cap_paragraph(doc, full_text, cap_size=42, cap_color=CLR_ACCENT,
                       body_size=10, font_body=FONT_BODY, font_cap=FONT_DISPL):
    """
    Simulasi drop cap (huruf pertama besar, sisa teks normal).
    python-docx tidak punya true drop-cap, jadi kita pakai huruf besar inline.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(4)
    first = full_text[0]
    rest  = full_text[1:]
    r_cap = p.add_run(first)
    _set_font(r_cap, name=font_cap, size=cap_size, bold=True, color=cap_color)
    r_rest = p.add_run(rest)
    _set_font(r_rest, name=font_body, size=body_size, color=CLR_DARK)
    return p

def pull_quote(doc, text, author=None):
    """Pull-quote besar di tengah, dibingkai garis atas-bawah merah."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    para_border(p,
                top={'color': CLR_ACCENT, 'sz': 18, 'space': '8'},
                bottom={'color': CLR_ACCENT, 'sz': 18, 'space': '8'})
    r1 = p.add_run('“ ')
    _set_font(r1, name=FONT_DISPL, size=22, bold=True, color=CLR_ACCENT)
    r2 = p.add_run(text)
    _set_font(r2, name=FONT_DISPL, size=13, italic=True, color=CLR_PRIMARY)
    r3 = p.add_run(' ”')
    _set_font(r3, name=FONT_DISPL, size=22, bold=True, color=CLR_ACCENT)
    if author:
        pa = doc.add_paragraph()
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pa.paragraph_format.space_before = Pt(0)
        pa.paragraph_format.space_after  = Pt(8)
        ra = pa.add_run(f'— {author}')
        _set_font(ra, name=FONT_DISPL, size=9, bold=True, color=CLR_ACCENT, small_caps=True)

def kicker(doc, text, color=CLR_ACCENT):
    """Kicker — teks kecil uppercase di atas headline (ala majalah)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text.upper())
    _set_font(r, name=FONT_DISPL, size=9, bold=True, color=color, small_caps=False)
    return p

def headline(doc, text, size=20, color=CLR_PRIMARY):
    """Headline artikel di dalam bab."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    _set_font(r, name=FONT_HEAD, size=size, bold=True, color=color)
    return p

def subhead(doc, text, size=12, color=CLR_PRIMARY):
    """Subheading dalam artikel."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    para_border(p, left={'color': CLR_ACCENT, 'sz': 18, 'space': '6'})
    p.paragraph_format.left_indent = Cm(0.25)
    r = p.add_run(text)
    _set_font(r, name=FONT_HEAD, size=size, bold=True, color=color)
    return p

def byline(doc, nama, meta=None):
    """Byline — Oleh: nama | meta"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    r1 = p.add_run('Oleh  ')
    _set_font(r1, name=FONT_DISPL, size=9, color=CLR_MUTE, small_caps=True)
    r2 = p.add_run(nama.upper())
    _set_font(r2, name=FONT_DISPL, size=9, bold=True, color=CLR_PRIMARY, small_caps=True)
    if meta:
        r3 = p.add_run(f'   •   {meta}')
        _set_font(r3, name=FONT_DISPL, size=9, color=CLR_MUTE)
    return p

def sidebar_box(doc, judul, items, bg=CLR_SIDEBAR, accent=CLR_SIDEACC,
                width_cm=None, font_size=9):
    """
    Kotak sidebar (mis. BerAKHLAK values) — tabel 1 kolom dengan heading & items.
    items: list of (label, isi) atau list of str.
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if width_cm:
        for cell in table.columns[0].cells:
            cell.width = Cm(width_cm)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)

    # judul sidebar
    p_jud = cell.paragraphs[0]
    p_jud.paragraph_format.space_before = Pt(0)
    p_jud.paragraph_format.space_after  = Pt(4)
    para_border(p_jud, bottom={'color': accent, 'sz': 8, 'space': '4'})
    r = p_jud.add_run(judul.upper())
    _set_font(r, name=FONT_HEAD, size=10, bold=True, color=accent, small_caps=True)

    for it in items:
        p_i = cell.add_paragraph()
        p_i.paragraph_format.space_before = Pt(2)
        p_i.paragraph_format.space_after  = Pt(2)
        p_i.paragraph_format.line_spacing = 1.2
        if isinstance(it, tuple):
            label, isi = it
            r_l = p_i.add_run('▪ ')
            _set_font(r_l, name=FONT_HEAD, size=font_size, bold=True, color=accent)
            r_l2 = p_i.add_run(f'{label}. ')
            _set_font(r_l2, name=FONT_HEAD, size=font_size, bold=True, color=accent)
            r_i = p_i.add_run(isi)
            _set_font(r_i, name=FONT_BODY, size=font_size, color=CLR_DARK)
        else:
            r_l = p_i.add_run('▪ ')
            _set_font(r_l, name=FONT_HEAD, size=font_size, bold=True, color=accent)
            r_i = p_i.add_run(str(it))
            _set_font(r_i, name=FONT_BODY, size=font_size, color=CLR_DARK)
    return table

def stat_card(cell, label, value, unit='', bg=CLR_PRIMARY, fg=CLR_PAPER):
    """Kartu statistik (angka besar)."""
    set_cell_bg(cell, bg)
    set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run(str(value))
    _set_font(r1, name=FONT_HEAD, size=26, bold=True, color=fg)
    if unit:
        r1u = p1.add_run(unit)
        _set_font(r1u, name=FONT_HEAD, size=12, bold=True, color=fg)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(0)
    r2 = p2.add_run(label.upper())
    _set_font(r2, name=FONT_DISPL, size=8, bold=True, color=fg, small_caps=True)

def infobox_table(doc, rows, header=None, bg_header=CLR_PRIMARY, bg_alt=CLR_SOFTGRAY):
    """Tabel ringkas ala infografis (striped)."""
    ncol = len(rows[0]) if not header else len(header)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if header:
        hr = t.add_row()
        for i, h in enumerate(header):
            c = hr.cells[i]
            set_cell_bg(c, bg_header)
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(h)
            _set_font(r, name=FONT_HEAD, size=9, bold=True, color=CLR_PAPER)
    for idx, row in enumerate(rows):
        tr = t.add_row()
        for i, val in enumerate(row):
            c = tr.cells[i]
            set_cell_margins(c, top=60, bottom=60, left=100, right=100)
            if idx % 2 == 1:
                set_cell_bg(c, bg_alt)
            c.text = ''
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.line_spacing = 1.2
            r = p.add_run(str(val))
            _set_font(r, name=FONT_BODY, size=9, bold=(i == 0), color=CLR_DARK)
    return t

def caption(doc, text):
    """Caption untuk gambar / grafik."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    _set_font(r, name=FONT_BODY, size=8, italic=True, color=CLR_MUTE)

def img_placeholder(doc, nomor, judul, deskripsi, tinggi_cm=5.5):
    """Placeholder gambar — kotak untuk sisipkan foto nanti."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, CLR_SOFTGRAY)
    set_cell_margins(cell, top=200, bottom=200, left=160, right=160)
    cell.height = Cm(tinggi_cm)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f'▢  GAMBAR {nomor}  ▢\n')
    _set_font(r1, name=FONT_HEAD, size=9, bold=True, color=CLR_PRIMARY)
    r2 = p.add_run(judul)
    _set_font(r2, name=FONT_BODY, size=9, italic=True, color=CLR_MUTE)
    caption(doc, f'Gambar {nomor}. {deskripsi}')

def chapter_cover(doc, nomor_bab, judul_bab, kicker_text, bg=CLR_PRIMARY, aksen=CLR_GOLD):
    """
    Cover pembuka bab — 1 halaman penuh dengan warna background.
    Dipasang di section 1-kolom dulu, lalu setelahnya section 2-kolom untuk isi.
    """
    # Tabel full-width sebagai 'kanvas'
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    set_cell_margins(cell, top=1800, bottom=1800, left=600, right=600)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Kicker atas
    p_k = cell.paragraphs[0]
    p_k.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_k.paragraph_format.space_before = Pt(0)
    p_k.paragraph_format.space_after  = Pt(6)
    r_k = p_k.add_run(kicker_text.upper())
    _set_font(r_k, name=FONT_DISPL, size=10, bold=True, color=aksen, small_caps=True)

    # Nomor bab besar
    p_no = cell.add_paragraph()
    p_no.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_no.paragraph_format.space_before = Pt(0)
    p_no.paragraph_format.space_after  = Pt(0)
    r_no = p_no.add_run(nomor_bab)
    _set_font(r_no, name=FONT_DISPL, size=72, bold=True, color=aksen)

    # Separator tipis
    p_sep = cell.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep.paragraph_format.space_before = Pt(6)
    p_sep.paragraph_format.space_after  = Pt(6)
    r_sep = p_sep.add_run('— • —')
    _set_font(r_sep, name=FONT_DISPL, size=14, bold=True, color=aksen)

    # Judul bab
    p_j = cell.add_paragraph()
    p_j.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_j.paragraph_format.space_before = Pt(0)
    p_j.paragraph_format.space_after  = Pt(12)
    p_j.paragraph_format.line_spacing = 1.15
    r_j = p_j.add_run(judul_bab)
    _set_font(r_j, name=FONT_DISPL, size=30, bold=True, color=CLR_PAPER)

    # Tagline bawah
    p_t = cell.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t.paragraph_format.space_before = Pt(8)
    p_t.paragraph_format.space_after  = Pt(0)
    r_t = p_t.add_run('MAJALAH AKTUALISASI  ·  SITRIA 2026')
    _set_font(r_t, name=FONT_DISPL, size=9, bold=True, color=aksen, small_caps=True)

    add_page_break(doc)

def add_running_header(section, teks_kiri='SITRIA', teks_kanan='MAJALAH AKTUALISASI 2026'):
    """Header halaman (masthead kecil)."""
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    # Tab stop kanan
    pPr = p._p.get_or_add_pPr()
    tabs = _oxml('w:tabs')
    tab = _oxml('w:tab', **{'w:val': 'right', 'w:leader': 'none', 'w:pos': '9000'})
    tabs.append(tab)
    pPr.append(tabs)
    para_border(p, bottom={'color': CLR_PRIMARY, 'sz': 4, 'space': '2'})
    r1 = p.add_run(teks_kiri)
    _set_font(r1, name=FONT_DISPL, size=9, bold=True, color=CLR_PRIMARY, small_caps=True)
    p.add_run('\t')
    r2 = p.add_run(teks_kanan)
    _set_font(r2, name=FONT_DISPL, size=8, color=CLR_MUTE, small_caps=True)

def add_page_footer(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_border(p, top={'color': CLR_PRIMARY, 'sz': 4, 'space': '2'})

    def _fld(para, instr):
        run = para.add_run()
        _set_font(run, name=FONT_DISPL, size=9, bold=True, color=CLR_PRIMARY)
        fc_begin = _oxml('w:fldChar', **{'w:fldCharType': 'begin'})
        instr_el = _oxml('w:instrText', **{'xml:space': 'preserve'})
        instr_el.text = f' {instr} '
        fc_end = _oxml('w:fldChar', **{'w:fldCharType': 'end'})
        run._r.append(fc_begin)
        run._r.append(instr_el)
        run._r.append(fc_end)

    r_left = p.add_run('— ')
    _set_font(r_left, name=FONT_DISPL, size=9, color=CLR_MUTE)
    _fld(p, 'PAGE')
    r_mid = p.add_run(' / ')
    _set_font(r_mid, name=FONT_DISPL, size=9, color=CLR_MUTE)
    _fld(p, 'NUMPAGES')
    r_right = p.add_run(' —')
    _set_font(r_right, name=FONT_DISPL, size=9, color=CLR_MUTE)


# ─────────────────────────────────────────────
#  DATA
# ─────────────────────────────────────────────

DATA = {
    'nama': 'Aidil Saputra Kirsan, S.ST., M.Tr.Kom',
    'nip': '199403172025061004',
    'jabatan': 'Dosen Asisten Ahli · Kepala Lab Inovasi Digital',
    'unit_kerja': 'Prodi Sistem Informasi — FSTI ITK',
    'instansi': 'Institut Teknologi Kalimantan',
    'mentor': 'Irma Fitria, S.Si., M.Si',
    'coach': 'Mustari Kurniawati, S.IP., MPA',
    'judul': 'SITRIA — Sistem Informasi Tridharma Akademik',
    'periode': '7 Maret – 22 April 2026',
    'tanggal_laporan': '15 April 2026',
    'tahun': '2026',
    'angkatan': 'VI',
}

NILAI_BERAKHLAK = [
    ('B', 'Berorientasi Pelayanan'),
    ('A', 'Akuntabel'),
    ('K', 'Kompeten'),
    ('H', 'Harmonis'),
    ('L', 'Loyal'),
    ('A', 'Adaptif'),
    ('K', 'Kolaboratif'),
]

# ─────────────────────────────────────────────
#  DOCUMENT SETUP
# ─────────────────────────────────────────────

doc = Document()

# Base style
style = doc.styles['Normal']
style.font.name = FONT_BODY
style.font.size = Pt(10)

# Section 0: halaman cover majalah — 1 kolom, tanpa header/footer
sec0 = doc.sections[0]
sec0.top_margin    = Cm(2.0)
sec0.bottom_margin = Cm(2.0)
sec0.left_margin   = Cm(2.0)
sec0.right_margin  = Cm(2.0)
sec0.different_first_page_header_footer = True


# ═══════════════════════════════════════════════
#  1. COVER MAJALAH
# ═══════════════════════════════════════════════

# Band atas — masthead
p_m = para(doc, '', space_before=0, space_after=2)
shade_para(p_m, CLR_PRIMARY)
para_border(p_m,
            top={'color': CLR_GOLD, 'sz': 24, 'space': '4'},
            bottom={'color': CLR_GOLD, 'sz': 24, 'space': '4'})
r_mast = p_m.add_run('  MAJALAH AKTUALISASI  ·  EDISI PERDANA  ·  APRIL 2026  ')
_set_font(r_mast, name=FONT_DISPL, size=10, bold=True, color=CLR_PAPER, small_caps=True)

# Judul besar
para(doc, '', space_before=24, space_after=0)
p_t = para(doc, 'SITRIA', font=FONT_DISPL, size=76, bold=True,
           color=CLR_PRIMARY, align='center', space_before=0, space_after=0,
           line_spacing=1.0)
p_sub = para(doc, 'Sistem Informasi Tridharma Akademik',
             font=FONT_DISPL, size=20, bold=True, color=CLR_ACCENT,
             align='center', space_before=0, space_after=4)
p_tag = para(doc,
             'Satu Dashboard untuk Riset, Publikasi, Pengabdian, dan Akreditasi',
             font=FONT_BODY, size=12, italic=True, color=CLR_DARK,
             align='center', space_after=18)

# Box hero teaser
tbl_hero = doc.add_table(rows=1, cols=1)
tbl_hero.alignment = WD_TABLE_ALIGNMENT.CENTER
c_hero = tbl_hero.rows[0].cells[0]
set_cell_bg(c_hero, CLR_CREAM)
set_cell_margins(c_hero, top=240, bottom=240, left=400, right=400)

p_h1 = c_hero.paragraphs[0]
p_h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_h1.paragraph_format.space_after = Pt(4)
r_h1 = p_h1.add_run('DARI LABORATORIUM KE DASHBOARD')
_set_font(r_h1, name=FONT_DISPL, size=11, bold=True, color=CLR_ACCENT, small_caps=True)

p_h2 = c_hero.add_paragraph()
p_h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_h2.paragraph_format.line_spacing = 1.25
p_h2.paragraph_format.space_after = Pt(6)
r_h2 = p_h2.add_run('Bagaimana seorang CPNS membangun sistem analitik Tridharma untuk '
                    'dua program studi — dari nol, dalam tujuh pekan.')
_set_font(r_h2, name=FONT_BODY, size=12, italic=True, color=CLR_PRIMARY)

p_h3 = c_hero.add_paragraph()
p_h3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_h3.paragraph_format.space_before = Pt(6)
p_h3.paragraph_format.space_after = Pt(0)
r_h3 = p_h3.add_run('Halaman 02 — 24')
_set_font(r_h3, name=FONT_DISPL, size=9, bold=True, color=CLR_MUTE, small_caps=True)

para(doc, '', space_before=24, space_after=0)

# Statbar hero
tbl_stat = doc.add_table(rows=1, cols=4)
tbl_stat.alignment = WD_TABLE_ALIGNMENT.CENTER
stat_card(tbl_stat.rows[0].cells[0], 'Dosen Terlayani', '25')
stat_card(tbl_stat.rows[0].cells[1], 'Fitur Dashboard', '7',  bg=CLR_ACCENT)
stat_card(tbl_stat.rows[0].cells[2], 'Program Studi',  '2',  bg=CLR_STEEL)
stat_card(tbl_stat.rows[0].cells[3], 'Pekan Aktualisasi', '7', bg=CLR_GOLD, fg=CLR_PRIMARY)

para(doc, '', space_before=20, space_after=0)

# Credit bawah cover
para(doc, 'Penyusun', font=FONT_DISPL, size=9, color=CLR_MUTE,
     align='center', space_after=2)
para(doc, DATA['nama'], font=FONT_DISPL, size=13, bold=True, color=CLR_PRIMARY,
     align='center', space_after=2)
para(doc, f"NIP. {DATA['nip']}  ·  {DATA['jabatan']}",
     font=FONT_BODY, size=9, italic=True, color=CLR_MUTE,
     align='center', space_after=2)
para(doc, DATA['unit_kerja'],
     font=FONT_BODY, size=9, italic=True, color=CLR_MUTE,
     align='center', space_after=16)

# Footer cover — band
p_b = para(doc, '', space_before=0, space_after=0)
shade_para(p_b, CLR_PRIMARY)
para_border(p_b,
            top={'color': CLR_GOLD, 'sz': 18, 'space': '3'},
            bottom={'color': CLR_GOLD, 'sz': 18, 'space': '3'})
r_b = p_b.add_run(f'  PELATIHAN DASAR CPNS  ·  ANGKATAN {DATA["angkatan"]}  ·  LAN RI  ·  {DATA["tahun"]}  ')
_set_font(r_b, name=FONT_DISPL, size=9, bold=True, color=CLR_PAPER, small_caps=True)

add_page_break(doc)


# ═══════════════════════════════════════════════
#  2. TABLE OF CONTENTS — GAYA MAJALAH
# ═══════════════════════════════════════════════

kicker(doc, 'Edisi Perdana · April 2026')
headline(doc, 'Daftar Isi', size=36, color=CLR_PRIMARY)
insert_hr_line(doc, color=CLR_ACCENT, sz=18)

# Feature items
toc_items = [
    ('BAB I',   'Rancangan Aktualisasi',
     'Profil instansi, isu terpilih, dan gagasan SITRIA.', '03'),
    ('BAB II',  'Implementasi Aktualisasi',
     'Lima kegiatan habituasi — dari pemetaan data hingga peluncuran sistem.', '08'),
    ('—',       'Kebermanfaatan Aktualisasi',
     'Dampak SITRIA bagi dosen, pimpinan, dan institusi.', '18'),
    ('BAB III', 'Penutup',
     'Kesimpulan dan rencana tindak lanjut keberlanjutan SITRIA.', '22'),
]

for bab, judul, teaser, hal in toc_items:
    # Row untuk 1 item
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c1, c2, c3 = t.rows[0].cells
    c1.width = Cm(2.8)
    c2.width = Cm(11.5)
    c3.width = Cm(1.8)
    set_cell_margins(c1, top=60, bottom=60, left=80, right=80)
    set_cell_margins(c2, top=60, bottom=60, left=120, right=80)
    set_cell_margins(c3, top=60, bottom=60, left=80, right=80)

    # Kolom 1: nomor bab
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p1.add_run(bab)
    _set_font(r1, name=FONT_DISPL, size=12, bold=True, color=CLR_ACCENT, small_caps=True)

    # Kolom 2: judul + teaser
    p2a = c2.paragraphs[0]
    p2a.paragraph_format.space_after = Pt(1)
    r2a = p2a.add_run(judul)
    _set_font(r2a, name=FONT_HEAD, size=14, bold=True, color=CLR_PRIMARY)
    p2b = c2.add_paragraph()
    p2b.paragraph_format.space_before = Pt(0)
    p2b.paragraph_format.space_after  = Pt(0)
    r2b = p2b.add_run(teaser)
    _set_font(r2b, name=FONT_BODY, size=9, italic=True, color=CLR_MUTE)

    # Kolom 3: nomor halaman
    p3 = c3.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = p3.add_run(hal)
    _set_font(r3, name=FONT_DISPL, size=20, bold=True, color=CLR_ACCENT)

    # Separator tipis
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(2)
    p_sep.paragraph_format.space_after  = Pt(2)
    para_border(p_sep, bottom={'color': CLR_SOFTGRAY, 'sz': 6, 'space': '2'})

# Editor's note
para(doc, '', space_before=12, space_after=0)
tbl_en = doc.add_table(rows=1, cols=1)
c_en = tbl_en.rows[0].cells[0]
set_cell_bg(c_en, CLR_CREAM)
set_cell_margins(c_en, top=200, bottom=200, left=260, right=260)

p_en1 = c_en.paragraphs[0]
p_en1.paragraph_format.space_after = Pt(4)
r_en1 = p_en1.add_run('CATATAN PENYUSUN')
_set_font(r_en1, name=FONT_DISPL, size=10, bold=True, color=CLR_ACCENT, small_caps=True)
p_en2 = c_en.add_paragraph()
p_en2.paragraph_format.line_spacing = 1.3
p_en2.paragraph_format.space_after  = Pt(0)
r_en2 = p_en2.add_run(
    'Majalah ini merangkum perjalanan habituasi saya selama tujuh pekan di Lab Inovasi Digital '
    'FSTI ITK. Dituliskan dalam format majalah agar lebih mudah dibaca oleh kolega, pimpinan, '
    'dan siapa pun yang tertarik melihat bagaimana nilai BerAKHLAK bertemu dengan kerja teknis '
    'sehari-hari. Selamat membaca.'
)
_set_font(r_en2, name=FONT_BODY, size=10, italic=True, color=CLR_DARK)
p_en3 = c_en.add_paragraph()
p_en3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_en3.paragraph_format.space_before = Pt(8)
p_en3.paragraph_format.space_after = Pt(0)
r_en3 = p_en3.add_run(f'— {DATA["nama"].split(",")[0]}')
_set_font(r_en3, name=FONT_DISPL, size=10, bold=True, color=CLR_PRIMARY, small_caps=True)

add_page_break(doc)


# ═══════════════════════════════════════════════
#  SECTION BREAK — HALAMAN ISI MULAI PAKAI HEADER/FOOTER
# ═══════════════════════════════════════════════
sec_isi = doc.add_section(WD_SECTION.NEW_PAGE)
sec_isi.top_margin    = Cm(2.0)
sec_isi.bottom_margin = Cm(2.0)
sec_isi.left_margin   = Cm(1.8)
sec_isi.right_margin  = Cm(1.8)
add_running_header(sec_isi)
add_page_footer(sec_isi)
set_columns(sec_isi, num=1)  # cover bab 1 kolom


# ═══════════════════════════════════════════════
#  BAB I — COVER BAB
# ═══════════════════════════════════════════════

chapter_cover(doc,
              nomor_bab='01',
              judul_bab='RANCANGAN\nAKTUALISASI',
              kicker_text='Bab Satu · Latar, Isu, dan Gagasan',
              bg=CLR_PRIMARY, aksen=CLR_GOLD)

# Pindah ke layout 2 kolom
sec_bab1 = doc.add_section(WD_SECTION.CONTINUOUS)
sec_bab1.top_margin    = Cm(2.0)
sec_bab1.bottom_margin = Cm(2.0)
sec_bab1.left_margin   = Cm(1.8)
sec_bab1.right_margin  = Cm(1.8)
add_running_header(sec_bab1, teks_kiri='BAB I · RANCANGAN')
add_page_footer(sec_bab1)
set_columns(sec_bab1, num=2, space_cm=0.7, sep=False)


# ── Artikel 1: Profil Instansi & Jabatan ──
kicker(doc, 'Tentang Institusi')
headline(doc, 'ITK: Institut Teknologi di Tepi Timur Indonesia', size=18)
byline(doc, DATA['nama'].split(',')[0], meta='Lab Inovasi Digital, FSTI ITK')

drop_cap_paragraph(doc,
    'Institut Teknologi Kalimantan (ITK) adalah perguruan tinggi negeri berbasis teknologi yang '
    'merupakan satu-satunya Institut Teknologi Negeri di wilayah Tengah dan Timur Indonesia. ITK '
    'didirikan berdasarkan Peraturan Presiden Nomor 125 Tahun 2014 dan diresmikan pada 6 Oktober 2014, '
    'berlokasi di Jalan Soekarno-Hatta KM 15, Balikpapan, Kalimantan Timur, di atas lahan seluas 300 hektare.')

body_para(doc,
    'ITK saat ini dipimpin oleh Rektor Prof. Dr. rer. nat. Agus Rubiyanto, M.Eng.Sc. '
    '(periode 2022–2026) dan telah mendapatkan akreditasi B dari BAN-PT. '
    'Fokus riset meliputi energi, kota cerdas, kemaritiman, dan teknologi pangan — sejalan dengan '
    'kebutuhan pembangunan Indonesia Tengah dan Timur.')

subhead(doc, 'Siapa Penulis?')
body_para(doc,
    'Penulis bertugas sebagai Dosen Program Studi Sistem Informasi sekaligus Kepala Laboratorium '
    'Inovasi Digital di Fakultas Sains dan Teknologi Industri (FSTI) ITK. Laboratorium ini '
    'merupakan unit penunjang akademik yang bertanggung jawab atas pengelolaan fasilitas, '
    'dukungan praktikum, serta pengembangan inovasi berbasis teknologi informasi.')

# Sidebar profil singkat
sidebar_box(doc, 'Snapshot Jabatan', [
    ('Jabatan', 'Dosen Asisten Ahli'),
    ('Peran Tambahan', 'Kepala Lab Inovasi Digital'),
    ('Unit Kerja', 'Prodi SI — FSTI ITK'),
    ('Pangkat', 'Penata Muda Tk.I / III-b'),
    ('Sistem Dikelola', 'SITRIA Dashboard Tridharma'),
])

# ── Artikel 2: Isu Terpilih ──
pull_quote(doc,
    'Data riset tersebar di tiga platform. Setiap kali akreditasi mendekat, rekap manual '
    'menghabiskan hari demi hari.',
    author='Latar Belakang Isu')

kicker(doc, 'Identifikasi Isu')
headline(doc, 'Tiga Isu, Satu yang Dipilih', size=18)

body_para(doc,
    'Berdasarkan pengamatan dan wawancara dengan pimpinan FSTI serta Kaprodi Sistem Informasi '
    'di lingkup Lab Inovasi Digital, teridentifikasi tiga isu aktual. Ketiga isu dianalisis '
    'menggunakan metode USG (Urgency, Seriousness, Growth), dan skor tertinggi menjadi fokus aktualisasi.')

infobox_table(doc,
    header=['Isu Aktual', 'U', 'S', 'G', 'Total', 'Prioritas'],
    rows=[
        ['Tugas Akhir belum terstandar',                4, 4, 3, 11, 'II'],
        ['Data riset & pengabdian belum terpusat ✓',    5, 5, 4, 14, 'I (Terpilih)'],
        ['Umpan balik mahasiswa belum sistematis',      3, 3, 3, 9,  'III'],
    ])

body_para(doc,
    'Isu terpilih — pengelolaan data riset dan pengabdian dosen yang belum terpusat dan teranalisis '
    '— mendapatkan skor USG = 14. Urgency tinggi (5) karena jadwal akreditasi mendesak; '
    'seriousness tinggi (5) karena berdampak langsung pada nilai akreditasi; growth (4) karena '
    'situasi semakin memburuk seiring bertambahnya jumlah dosen dan publikasi.')


# ── Artikel 3: Gagasan Kreatif SITRIA ──
pull_quote(doc,
    'SITRIA bukan sekadar dashboard — ia adalah jembatan antara data mentah dan keputusan strategis.',
    author='Visi Sistem')

kicker(doc, 'Gagasan Kreatif')
headline(doc, 'Lahirnya SITRIA — Dashboard untuk 25 Dosen & Dua Prodi', size=18)

body_para(doc,
    'Berdasarkan isu terpilih, gagasan yang diusulkan adalah pengembangan SITRIA (Sistem Informasi '
    'Tridharma Akademik): dashboard analitik berbasis web yang mengintegrasikan dan menganalisis '
    'data riset, publikasi, dan pengabdian seluruh dosen Prodi Sistem Informasi dan Prodi Bisnis '
    'Digital FSTI ITK secara otomatis dari portal SINTA Kemdiktisaintek.')

body_para(doc,
    'SITRIA dirancang dengan teknologi web modern dan dilengkapi kecerdasan buatan untuk pengelompokan '
    'topik riset otomatis. Tujuh fitur utama diramu agar dapat melayani dosen, Kaprodi, pimpinan, '
    'hingga mahasiswa dan mitra riset.')

# Sidebar 7 fitur
sidebar_box(doc, 'Tujuh Fitur SITRIA', [
    ('F1', 'Dashboard Analitik Utama — metrik Tridharma 2 prodi.'),
    ('F2', 'Research Gallery — 6 kategori karya akademik.'),
    ('F3', 'AI Clustering — deteksi topik & potensi kolaborasi.'),
    ('F4', 'Sankey Timeline — evolusi topik riset 2018→kini.'),
    ('F5', 'Funding Dashboard — monitoring dana & hibah BIMA.'),
    ('F6', 'DTPS Akreditasi — kalkulasi rasio LKPS otomatis.'),
    ('F7', 'Expertise Finder — pencocokan pakar berbasis AI.'),
])

# ── Artikel 4: Tujuan & Manfaat ──
kicker(doc, 'Tujuan Aktualisasi')
headline(doc, 'Tiga Tujuan, Empat Penerima Manfaat', size=18)

body_para(doc, 'Tujuan aktualisasi disusun dalam tiga poin utama:')

for i, t in enumerate([
    'Memperkuat karakter ASN penulis melalui internalisasi nilai-nilai BerAKHLAK dan penerapan '
    'prinsip Smart ASN serta Manajemen ASN di lingkungan kerja.',
    'Menganalisis dan mencari alternatif solusi terhadap isu pengelolaan data riset dan pengabdian '
    'dosen di FSTI ITK yang belum terpusat.',
    'Membangun dashboard analitik Tridharma terpusat (SITRIA) yang mengintegrasikan data riset, '
    'publikasi, dan pendanaan dosen secara otomatis.',
], 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.line_spacing = 1.25
    r1 = p.add_run(f'{i}.  ')
    _set_font(r1, name=FONT_HEAD, size=10, bold=True, color=CLR_ACCENT)
    r2 = p.add_run(t)
    _set_font(r2, name=FONT_BODY, size=10, color=CLR_DARK)

subhead(doc, 'Siapa yang Diuntungkan?')
infobox_table(doc,
    header=['Pihak', 'Manfaat Utama'],
    rows=[
        ['Penulis (ASN Dosen)',
         'Kompetensi teknis & manajerial; aktualisasi BerAKHLAK nyata; kapasitas Smart ASN.'],
        ['Lab & Prodi SI',
         'Data riset terpusat; keputusan berbasis data; persiapan akreditasi lebih mudah.'],
        ['Institusi ITK',
         'Akuntabilitas kinerja penelitian; penguatan akreditasi; model best practice.'],
        ['Ekosistem Riset',
         'Peluang kolaborasi lintas prodi; peta riset ITK; produktivitas kolektif.'],
    ])

# ── Artikel 5: Matriks Rencana ──
pull_quote(doc,
    'Lima kegiatan, tujuh pekan, tujuh nilai BerAKHLAK yang dihabituasikan di setiap langkah.',
    author='Rencana Aktualisasi')

kicker(doc, 'Matriks Kegiatan')
headline(doc, 'Peta Jalan Tujuh Pekan', size=18)

infobox_table(doc,
    header=['#', 'Kegiatan', 'Periode', 'Output'],
    rows=[
        ['1', 'Pemetaan & Pengumpulan Data Dosen', '07–20 Mar',
         'Basis data 25 dosen terstruktur'],
        ['2', 'Sistem Pengambilan Data SINTA Otomatis', '14–27 Mar',
         'Scraper aktif + penjadwalan berkala'],
        ['3', 'Fitur Analitik & Dashboard SITRIA', '21 Mar – 10 Apr',
         '7 fitur dashboard berfungsi penuh'],
        ['4', 'Dashboard DTPS & Pendanaan Riset', '28 Mar – 17 Apr',
         'DTPS & Funding aktif + berita acara'],
        ['5', 'Peluncuran, Sosialisasi & Evaluasi', '11–22 Apr',
         'Sistem online, SOP, laporan evaluasi'],
    ])

body_para(doc,
    'Setiap kegiatan dirancang agar seluruh tujuh nilai BerAKHLAK — Berorientasi Pelayanan, Akuntabel, '
    'Kompeten, Harmonis, Loyal, Adaptif, Kolaboratif — terhabitusi dalam setiap langkah teknis '
    'maupun interaksi dengan pemangku kepentingan.')


# ═══════════════════════════════════════════════
#  BAB II — COVER BAB
# ═══════════════════════════════════════════════

# Kembali ke 1 kolom untuk cover bab
sec_bab2cov = doc.add_section(WD_SECTION.NEW_PAGE)
sec_bab2cov.top_margin    = Cm(2.0)
sec_bab2cov.bottom_margin = Cm(2.0)
sec_bab2cov.left_margin   = Cm(1.8)
sec_bab2cov.right_margin  = Cm(1.8)
add_running_header(sec_bab2cov, teks_kiri='BAB II · IMPLEMENTASI')
add_page_footer(sec_bab2cov)
set_columns(sec_bab2cov, num=1)

chapter_cover(doc,
              nomor_bab='02',
              judul_bab='IMPLEMENTASI\nAKTUALISASI',
              kicker_text='Bab Dua · Lima Kegiatan, Tujuh Pekan',
              bg=CLR_ACCENT, aksen=CLR_GOLD)

# Pindah ke 2 kolom untuk isi
sec_bab2 = doc.add_section(WD_SECTION.CONTINUOUS)
sec_bab2.top_margin    = Cm(2.0)
sec_bab2.bottom_margin = Cm(2.0)
sec_bab2.left_margin   = Cm(1.8)
sec_bab2.right_margin  = Cm(1.8)
add_running_header(sec_bab2, teks_kiri='BAB II · IMPLEMENTASI')
add_page_footer(sec_bab2)
set_columns(sec_bab2, num=2, space_cm=0.7)


# ── Opening BAB II ──
kicker(doc, 'Ringkasan Pelaksanaan')
headline(doc, 'Dari Rencana ke Realisasi — 100% Tuntas', size=18)
byline(doc, DATA['nama'].split(',')[0], meta=f"Periode {DATA['periode']}")

drop_cap_paragraph(doc,
    'Selama periode 7 Maret hingga 22 April 2026, kelima kegiatan yang direncanakan berhasil '
    'diselesaikan seluruhnya. Setiap kegiatan dilaksanakan dengan pendekatan 5W1H, melibatkan '
    'pemangku kepentingan yang relevan, dan dilandasi nilai-nilai BerAKHLAK. Berikut catatan '
    'pelaksanaan dari kelima kegiatan tersebut.')


# ── Data kegiatan ──
KEGIATAN = [
    {
        'no': 1,
        'judul': 'Pemetaan & Pengumpulan Data Dosen',
        'subjudul': 'Fondasi data 25 dosen dari dua program studi',
        'periode': '07 – 20 Maret 2026 (Minggu 1–2)',
        'lead': ('Tanpa data yang valid, seluruh fitur analitik SITRIA tidak akan berfungsi. '
                 'Langkah pertama adalah memastikan setiap dosen Prodi SI (15 orang) dan Prodi '
                 'Bisnis Digital (10 orang) terpetakan lengkap, terverifikasi, dan siap dikonsumsi sistem.'),
        'tahapan': [
            'Koordinasi dengan Kaprodi SI & Bisnis Digital',
            'Pengumpulan SINTA ID seluruh dosen aktif',
            'Verifikasi via portal SINTA Kemdiktisaintek',
            'Penyusunan basis data terstruktur (JSON)',
            'Validasi bersama Mentor dan Kaprodi',
        ],
        'capaian': [
            'Data 25 dosen dari 2 prodi berhasil dikumpulkan lengkap.',
            'Seluruhnya terverifikasi via profil resmi SINTA.',
            'Tersusun dalam format JSON siap-pakai untuk fitur SITRIA.',
            'Validasi resmi diperoleh dari Koordinator Prodi.',
        ],
        'berakhlak_quote': 'Akurasi data adalah pelayanan — pondasi keputusan yang akuntabel.',
        'bukti': [
            ('1.1', 'Rapat koordinasi Kaprodi SI & Bisnis Digital'),
            ('1.2', 'Profil dosen di portal SINTA Kemdiktisaintek'),
            ('1.3', 'Basis data dosen terstruktur (lecturers.json)'),
            ('1.4', 'Lembar validasi data dosen (ttd Kaprodi)'),
        ],
    },
    {
        'no': 2,
        'judul': 'Sistem Pengambilan Data SINTA Otomatis',
        'subjudul': 'Dari jam-jaman menjadi menit-menitan',
        'periode': '14 – 27 Maret 2026 (Minggu 2–3)',
        'lead': ('Pengumpulan manual rawan salah dan menghabiskan waktu. Solusinya: membangun '
                 'sistem Python yang mengambil data riset, publikasi, dan pengabdian seluruh dosen '
                 'dari portal SINTA secara otomatis — dan menjadwalkannya.'),
        'tahapan': [
            'Analisis struktur data portal SINTA',
            'Pengembangan scraper Python',
            'Konfigurasi penjadwalan otomatis berkala',
            'Pengujian dengan data kedua prodi',
            'Penyusunan dokumentasi teknis',
        ],
        'capaian': [
            'Scraper Python aktif untuk 25 dosen dari dua prodi.',
            '4 kelompok data/prodi: statistik, cluster, keahlian, peta riset.',
            'Proses yang tadinya berjam-jam, kini selesai dalam hitungan menit.',
            'Penjadwalan otomatis terkonfigurasi, berjalan berkala.',
            'Dokumentasi teknis lengkap disusun untuk keberlanjutan.',
        ],
        'berakhlak_quote': 'Kompeten berarti memecahkan masalah dengan keahlian yang dimiliki.',
        'bukti': [
            ('2.1', 'Program pengambilan data SINTA sedang berjalan'),
            ('2.2', 'Data SINTA yang tersimpan di sistem'),
            ('2.3', 'Konfigurasi penjadwalan otomatis (Task Scheduler)'),
            ('2.4', 'Dokumentasi teknis sistem'),
        ],
    },
    {
        'no': 3,
        'judul': 'Fitur Analitik & Dashboard SITRIA',
        'subjudul': 'Lima fitur inti — dari AI clustering hingga expertise finder',
        'periode': '21 Maret – 10 April 2026 (Minggu 3–5)',
        'lead': ('Fitur analitik adalah jantung SITRIA — ia mengubah data mentah menjadi intelijen '
                 'strategis yang mudah dipahami pimpinan dan dosen. Lima dari tujuh fitur '
                 'diimplementasikan pada fase ini.'),
        'tahapan': [
            'Implementasi AI Clustering untuk topik riset',
            'Pengembangan visualisasi Sankey Timeline',
            'Pembangunan Research Gallery 6 kategori',
            'Integrasi frontend dengan pipeline data',
            'Pengujian menyeluruh seluruh modul',
        ],
        'capaian': [
            'AI Clustering mendeteksi kesamaan topik & potensi kolaborasi lintas prodi.',
            'Sankey Timeline menampilkan evolusi riset dari 2018 sampai kini.',
            'Research Gallery mendukung 6 kategori karya akademik.',
            'Expertise Finder memudahkan pencarian pakar riset.',
            'Seluruh komponen dashboard terhubung dengan pipeline data.',
        ],
        'berakhlak_quote': 'Adaptif bukan sekadar mengikuti tren — tapi menyalurkannya ke manfaat nyata.',
        'bukti': [
            ('3.1', 'Dashboard Analitik Utama SITRIA'),
            ('3.2', 'Halaman AI Clustering / Pengelompokan Topik'),
            ('3.3', 'Sankey Timeline perjalanan riset'),
            ('3.4', 'Galeri Karya Akademik 6 kategori'),
            ('3.5', 'Expertise Finder — pencarian pakar'),
        ],
    },
    {
        'no': 4,
        'judul': 'Dashboard DTPS & Pendanaan Riset',
        'subjudul': 'Akreditasi yang sebelumnya manual kini otomatis',
        'periode': '28 Maret – 17 April 2026 (Minggu 4–6)',
        'lead': ('Rasio DTPS adalah jantung LKPS BAN-PT. Sebelumnya ia disusun manual dari '
                 'spreadsheet terpisah. SITRIA menghadirkan dashboard yang mengalkulasi tiga rasio '
                 'tersebut otomatis, divalidasi langsung bersama Wakil Dekan Akademik.'),
        'tahapan': [
            'Kajian standar LKPS BAN-PT',
            'Implementasi kalkulasi otomatis tiga rasio DTPS',
            'Pengembangan Funding Dashboard (BIMA)',
            'Integrasi data hibah internal & eksternal',
            'Validasi dengan Wakil Dekan (13 Apr 2026)',
        ],
        'capaian': [
            'Rasio Penelitian/DTPS, Pengabdian/DTPS, Dana/DTPS otomatis.',
            'Fitur simulasi interaktif untuk skenario komposisi dosen.',
            'Dashboard Pendanaan menampilkan aliran hibah 2 prodi.',
            'Validasi resmi Wakil Dekan FSTI — dinyatakan valid.',
        ],
        'berakhlak_quote': 'Akuntabel berarti angka yang bisa diverifikasi — kapan saja, oleh siapa saja.',
        'bukti': [
            ('4.1', 'Dashboard DTPS Akreditasi'),
            ('4.2', 'Dashboard Pemantauan Dana & Hibah'),
            ('4.3', 'Simulasi interaktif DTPS'),
            ('4.4', 'Berita Acara Validasi DTPS (ttd Wakil Dekan)'),
        ],
    },
    {
        'no': 5,
        'judul': 'Peluncuran, Sosialisasi & Evaluasi',
        'subjudul': 'Dari prototipe ke sistem produksi — 23 peserta, rating 4,3/5,0',
        'periode': '11 – 22 April 2026 (Minggu 6–7)',
        'lead': ('Sistem yang hebat tanpa adopsi hanyalah kode yang sepi. Kegiatan kelima memastikan '
                 'SITRIA diluncurkan, dilatihkan, dievaluasi — lengkap dengan SOP dan panduan pengguna '
                 'untuk keberlanjutan pasca-aktualisasi.'),
        'tahapan': [
            'Peluncuran ke server Lab Inovasi Digital (15 Apr)',
            'Penyusunan Panduan Pengguna & SOP',
            'Sosialisasi & pelatihan (16–17 Apr)',
            'Pengumpulan evaluasi & masukan',
            'Penyusunan laporan akhir aktualisasi',
        ],
        'capaian': [
            'SITRIA aktif di server Lab Inovasi Digital (15 Apr 2026).',
            'Sosialisasi 23 peserta — dosen, Kaprodi, laboran.',
            'Panduan Pengguna & SOP didistribusikan.',
            'Rata-rata kepuasan pengguna 4,3/5,0.',
            'Seluruh 5 kegiatan 100% tuntas dalam 7 pekan.',
        ],
        'berakhlak_quote': 'Kolaboratif berarti melibatkan pengguna sejak hari pertama hingga launch day.',
        'bukti': [
            ('5.1', 'SITRIA live di server Lab Inovasi Digital'),
            ('5.2', 'Pelaksanaan sosialisasi (16 Apr)'),
            ('5.3', 'Sesi pelatihan penggunaan dashboard (17 Apr)'),
            ('5.4', 'Daftar hadir 23 peserta sosialisasi'),
            ('5.5', 'Rekap evaluasi kepuasan pengguna (4,3/5,0)'),
        ],
    },
]

for kg in KEGIATAN:
    # Separator halus antar kegiatan
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(10)
    p_sep.paragraph_format.space_after  = Pt(4)
    para_border(p_sep, top={'color': CLR_ACCENT, 'sz': 12, 'space': '4'})

    kicker(doc, f"Kegiatan {kg['no']} · {kg['periode']}")
    headline(doc, kg['judul'], size=18)
    para(doc, kg['subjudul'], font=FONT_BODY, size=11, italic=True, color=CLR_MUTE,
         space_before=0, space_after=6)

    # Lead paragraph dengan drop cap
    drop_cap_paragraph(doc, kg['lead'])

    # Pull quote BerAKHLAK
    pull_quote(doc, kg['berakhlak_quote'], author=f'Kegiatan {kg["no"]}')

    # Tahapan — sidebar
    sidebar_box(doc, 'Tahapan Pelaksanaan',
                [(str(i+1), t) for i, t in enumerate(kg['tahapan'])])

    # Capaian
    subhead(doc, 'Capaian Utama')
    for c in kg['capaian']:
        p_c = doc.add_paragraph()
        p_c.paragraph_format.space_before = Pt(1)
        p_c.paragraph_format.space_after  = Pt(2)
        p_c.paragraph_format.left_indent = Cm(0.5)
        p_c.paragraph_format.first_line_indent = Cm(-0.5)
        p_c.paragraph_format.line_spacing = 1.25
        r1 = p_c.add_run('▪  ')
        _set_font(r1, name=FONT_HEAD, size=10, bold=True, color=CLR_ACCENT)
        r2 = p_c.add_run(c)
        _set_font(r2, name=FONT_BODY, size=10, color=CLR_DARK)

    # Bukti — grid 2 kolom, pakai placeholder mini
    subhead(doc, 'Bukti Dokumentasi')
    for no, desc in kg['bukti']:
        img_placeholder(doc, no, desc, desc, tinggi_cm=4.5)


# ── 2.2 Kebermanfaatan Aktualisasi ──
p_sep = doc.add_paragraph()
p_sep.paragraph_format.space_before = Pt(14)
p_sep.paragraph_format.space_after  = Pt(4)
para_border(p_sep, top={'color': CLR_PRIMARY, 'sz': 18, 'space': '4'})

kicker(doc, 'Kebermanfaatan Aktualisasi')
headline(doc, 'Dampak Nyata bagi Banyak Pihak', size=18)

drop_cap_paragraph(doc,
    'Sistem SITRIA yang telah diluncurkan pada 15 April 2026 memberikan kebermanfaatan nyata bagi '
    'berbagai pemangku kepentingan di lingkungan Prodi Sistem Informasi dan Bisnis Digital FSTI ITK. '
    'Berdasarkan evaluasi pada 9–14 April 2026 dengan 23 responden, rata-rata kepuasan pengguna '
    'mencapai 4,3 dari skala 5,0.')

subhead(doc, 'Bagi Stakeholders')
infobox_table(doc,
    header=['Pihak', 'Manfaat', 'Indikator'],
    rows=[
        ['Dosen (25 SI+Bisdig)',
         'Pantau profil riset mandiri; rekomendasi kolaborasi dari AI.',
         '25 dosen aktif · 89% data akurat'],
        ['Kaprodi SI & Bisdig',
         'DTPS otomatis — tidak perlu rekap manual.',
         'Efisiensi 60–70% waktu penyusunan'],
        ['Pimpinan FSTI',
         'Monitoring real-time + simulasi komposisi dosen.',
         'Divalidasi Wadek 13 Apr 2026'],
        ['Mahasiswa & Mitra',
         'Expertise Finder — dari hari ke menit.',
         'Akses publik, proses jauh lebih cepat'],
        ['Laboran',
         'Pengelolaan data terstruktur + SOP jelas.',
         '1 laboran terlatih + SOP tersusun'],
    ])

subhead(doc, 'Bagi Organisasi / Instansi')
infobox_table(doc,
    header=['Dimensi', 'Manfaat bagi Prodi SI, Bisdig & FSTI ITK'],
    rows=[
        ['Kesiapan Akreditasi',
         'Tiga indikator LKPS real-time — prodi bisa ambil langkah korektif sebelum akreditasi.'],
        ['Transformasi Digital',
         'Akhir dari era rekap manual — 3 platform (SINTA, Scopus, Scholar) kini satu dashboard.'],
        ['Produktivitas Riset',
         'AI Clustering & Expertise Finder mendorong kolaborasi lintas prodi.'],
        ['Visibilitas Institusi',
         'Research Gallery publik + Sankey Timeline meningkatkan jejak riset institusi.'],
        ['Potensi Replikasi',
         'Arsitektur modular — dapat direplikasi ke prodi lain di ITK.'],
        ['Budaya BerAKHLAK',
         'Data-driven culture mulai tertanam — selaras Akuntabel, Kompeten, Adaptif.'],
    ])


# ═══════════════════════════════════════════════
#  BAB III — COVER BAB
# ═══════════════════════════════════════════════
sec_bab3cov = doc.add_section(WD_SECTION.NEW_PAGE)
sec_bab3cov.top_margin    = Cm(2.0)
sec_bab3cov.bottom_margin = Cm(2.0)
sec_bab3cov.left_margin   = Cm(1.8)
sec_bab3cov.right_margin  = Cm(1.8)
add_running_header(sec_bab3cov, teks_kiri='BAB III · PENUTUP')
add_page_footer(sec_bab3cov)
set_columns(sec_bab3cov, num=1)

chapter_cover(doc,
              nomor_bab='03',
              judul_bab='PENUTUP',
              kicker_text='Bab Tiga · Kesimpulan & Tindak Lanjut',
              bg=CLR_STEEL, aksen=CLR_GOLD)

# Pindah ke 2 kolom
sec_bab3 = doc.add_section(WD_SECTION.CONTINUOUS)
sec_bab3.top_margin    = Cm(2.0)
sec_bab3.bottom_margin = Cm(2.0)
sec_bab3.left_margin   = Cm(1.8)
sec_bab3.right_margin  = Cm(1.8)
add_running_header(sec_bab3, teks_kiri='BAB III · PENUTUP')
add_page_footer(sec_bab3)
set_columns(sec_bab3, num=2, space_cm=0.7)


# ── 3.1 Kesimpulan ──
kicker(doc, 'Kesimpulan')
headline(doc, 'Tujuh Pekan, Tujuh Nilai, Satu Sistem Hidup', size=18)

drop_cap_paragraph(doc,
    'Kegiatan aktualisasi SITRIA (Sistem Informasi Tridharma Akademik) dalam rangka Pelatihan Dasar '
    'CPNS Angkatan VI Tahun 2026 telah berhasil dilaksanakan selama periode 7 Maret – 22 April 2026 '
    'dengan tingkat penyelesaian 100%. Tiga kesimpulan utama dapat diambil dari perjalanan ini.')

kesimpulan_items = [
    ('Tercapainya Seluruh Kegiatan',
     'Kelima kegiatan berhasil diselesaikan 100% sesuai jadwal. SITRIA aktif di server Lab Inovasi '
     'Digital FSTI ITK dengan 7 fitur utama, melayani 25 dosen dari 2 program studi.'),
    ('Terhabitusikannya BerAKHLAK',
     'Seluruh rangkaian dijalankan dengan internalisasi BerAKHLAK nyata: Berorientasi Pelayanan '
     '(akses data mudah), Akuntabel (kalkulasi teraudit), Kompeten (keahlian teknis AI/web), '
     'Harmonis (kolaborasi dengan Kaprodi & pimpinan), Loyal (mendukung akreditasi), Adaptif '
     '(integrasi multi-sumber), dan Kolaboratif (melibatkan pemangku sejak awal).'),
    ('Terselesaikannya Isu',
     'Isu terpilih — data riset tersebar di SINTA, Scopus, Scholar — teratasi dalam satu dashboard '
     'terpusat. DTPS akreditasi yang dulu butuh berhari-hari kini hitungan jam. Rata-rata kepuasan '
     'pengguna: 4,3/5,0.'),
]

for judul_k, narasi_k in kesimpulan_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.line_spacing = 1.3
    r0 = p.add_run('▸  ')
    _set_font(r0, name=FONT_HEAD, size=11, bold=True, color=CLR_ACCENT)
    r1 = p.add_run(f'{judul_k}. ')
    _set_font(r1, name=FONT_HEAD, size=10, bold=True, color=CLR_PRIMARY)
    r2 = p.add_run(narasi_k)
    _set_font(r2, name=FONT_BODY, size=10, color=CLR_DARK)


# ── 3.2 Tindak Lanjut ──
pull_quote(doc,
    'Aktualisasi boleh berakhir — tapi sistem, kolaborasi, dan nilai BerAKHLAK baru saja dimulai.',
    author='Pesan Penutup')

kicker(doc, 'Tindak Lanjut')
headline(doc, 'Enam Langkah Keberlanjutan SITRIA', size=18)

body_para(doc,
    'Agar SITRIA dan budaya kerja berbasis data terus berkembang pasca-aktualisasi, enam rencana '
    'tindak lanjut berikut direkomendasikan kepada para pemangku kepentingan:')

infobox_table(doc,
    header=['#', 'Rencana Tindak Lanjut', 'Target'],
    rows=[
        ['1', 'Integrasi dengan SISTER & SIMPEG ITK — data kepegawaian real-time.',
         'Sem. I 2026/2027'],
        ['2', 'Fitur ekspor laporan otomatis (PDF/Excel) untuk borang akreditasi.',
         'Agustus 2026'],
        ['3', 'Replikasi SITRIA ke prodi lain di ITK — model best practice.',
         'TA 2026/2027'],
        ['4', 'Optimalisasi AI Clustering dengan data historis & algoritma lebih presisi.',
         'Oktober 2026'],
        ['5', 'SK Dekan/Rektor — SITRIA sebagai sistem resmi pengelolaan Tridharma FSTI.',
         'Sem. I 2026/2027'],
        ['6', 'Pelatihan berkala pengelola data prodi & laboran baru.',
         'Setiap awal sem.'],
    ])

# Closing paragraph
para(doc, '', space_before=8)
body_para(doc,
    'Demikian Laporan Aktualisasi ini disusun sebagai pertanggungjawaban atas pelaksanaan kegiatan '
    'habituasi Pelatihan Dasar CPNS Angkatan VI Tahun 2026. Penulis berharap SITRIA dapat terus '
    'berkembang dan memberikan manfaat yang berkelanjutan bagi Prodi Sistem Informasi, Prodi '
    'Bisnis Digital, FSTI, dan Institut Teknologi Kalimantan secara keseluruhan.')

# Signature
para(doc, '', space_before=16)
para(doc, f'Balikpapan, {DATA["tanggal_laporan"]}',
     font=FONT_BODY, size=10, align='right', space_after=4)
para(doc, 'Penyusun,', font=FONT_BODY, size=10, align='right', space_after=40)
p_nm = para(doc, DATA['nama'], font=FONT_HEAD, size=11, bold=True, color=CLR_PRIMARY,
            align='right', space_after=2)
para(doc, f"NIP. {DATA['nip']}", font=FONT_BODY, size=10, align='right', space_after=2)


# ═══════════════════════════════════════════════
#  BACK COVER — COLOPHON
# ═══════════════════════════════════════════════
sec_bc = doc.add_section(WD_SECTION.NEW_PAGE)
sec_bc.top_margin    = Cm(2.5)
sec_bc.bottom_margin = Cm(2.0)
sec_bc.left_margin   = Cm(2.5)
sec_bc.right_margin  = Cm(2.5)
set_columns(sec_bc, num=1)

para(doc, '', space_before=40)
para(doc, 'SITRIA', font=FONT_DISPL, size=48, bold=True, color=CLR_PRIMARY,
     align='center', space_after=0)
para(doc, 'MAJALAH AKTUALISASI · EDISI PERDANA',
     font=FONT_DISPL, size=11, bold=True, color=CLR_ACCENT,
     align='center', space_after=24)

insert_hr_line(doc, color=CLR_PRIMARY, sz=18)
para(doc, 'COLOPHON', font=FONT_DISPL, size=10, bold=True, color=CLR_ACCENT,
     align='center', space_before=6, space_after=6)

colophon_lines = [
    ('Penyusun', DATA['nama']),
    ('NIP', DATA['nip']),
    ('Unit Kerja', DATA['unit_kerja']),
    ('Mentor', DATA['mentor']),
    ('Coach', DATA['coach']),
    ('Judul Aktualisasi', DATA['judul']),
    ('Periode Habituasi', DATA['periode']),
    ('Tanggal Terbit', DATA['tanggal_laporan']),
]
for k, v in colophon_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f'{k}  ·  ')
    _set_font(r1, name=FONT_DISPL, size=9, color=CLR_MUTE, small_caps=True)
    r2 = p.add_run(v)
    _set_font(r2, name=FONT_BODY, size=10, bold=True, color=CLR_PRIMARY)

insert_hr_line(doc, color=CLR_PRIMARY, sz=18)

para(doc, 'Dicetak dan disusun di Balikpapan',
     font=FONT_BODY, size=9, italic=True, color=CLR_MUTE,
     align='center', space_before=8, space_after=2)
para(doc, 'Pelatihan Dasar CPNS · LAN RI · 2026',
     font=FONT_DISPL, size=9, bold=True, color=CLR_PRIMARY,
     align='center', space_after=2)


# ═══════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════
output_path = r'd:\Github-ADL\presentasi-umum\CPNS\laporan-progress\Majalah_Aktualisasi_SITRIA_2026.docx'
doc.save(output_path)
print(f'[OK] Majalah aktualisasi berhasil dibuat: {output_path}')
